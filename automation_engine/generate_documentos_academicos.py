"""
Generador de documentos academicos (ACA, PRESENTACION, FORO) a partir de
exactamente 5 archivos fuente analizados con OpenAI.

No depende de plantillas .docx externas: la estructura de cada documento
esta definida directamente en el prompt y en el codigo.

Reutiliza helpers de `generate_guiones.py` sin modificarlo.
"""

import argparse
import os
import re
import sys
import time
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

try:
    from dotenv import load_dotenv
except ImportError:  # pragma: no cover
    load_dotenv = None

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover
    OpenAI = None

from automation_engine.config.categories import resolve_docx_prompt
from automation_engine.generate_guiones import (
    clean_text,
    extract_docx_text,
    extract_pdf_text,
    load_system_prompt,
    word_count,
)


ENGINE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = ENGINE_DIR.parent
DEFAULT_PROMPT_PATH = PROJECT_ROOT / "prompts" / "docx" / "system_prompt_documentos_academicos.md"
DEFAULT_INPUT_DIR = PROJECT_ROOT / "inputs"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "outputs" / "documentos_academicos"
DEFAULT_SUBJECT = "Asignatura"
DEFAULT_PROGRAM = "Programa"

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
EXPECTED_FILE_COUNT = 5

DOCUMENT_TYPES: Tuple[str, ...] = ("ACA", "PRESENTACION", "FORO")
DOCUMENT_TITLES: Dict[str, str] = {
    "ACA": "PROYECTO FINAL (ACA)",
    "PRESENTACION": "PRESENTACION DE LA ASIGNATURA",
    "FORO": "FORO ACADEMICO",
}

DOCX_PROMPTS: Dict[str, str] = {
    "ACA": "prompts/docx/docx_aca.md",
    "PRESENTACION": "prompts/docx/docx_presentacion.md",
    "FORO": "prompts/docx/docx_foro.md",
}

FONT_NAME = "Arial"
BODY_FONT_SIZE = 12
HEADING_FONT_SIZE = 14
BLACK = RGBColor(0x00, 0x00, 0x00)

MAX_CONTENT_CHARS_PER_FILE = 25000


def extract_txt_text(path: Path) -> str:
    text = ""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
    if not text:
        text = path.read_text(encoding="utf-8", errors="ignore")
    return "\n".join(clean_text(line) for line in text.splitlines() if clean_text(line))


def extract_input_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".txt":
        return extract_txt_text(path)
    raise ValueError(f"Formato no soportado: {suffix}. Usa .docx, .pdf o .txt")


def collect_input_files(input_dir: Path) -> List[Path]:
    if not input_dir.exists() or not input_dir.is_dir():
        raise FileNotFoundError(
            f"No se encontro la carpeta de entrada: {input_dir}"
        )
    files = sorted(
        path
        for path in input_dir.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    if len(files) != EXPECTED_FILE_COUNT:
        raise ValueError(
            f"Se esperaban exactamente {EXPECTED_FILE_COUNT} archivos en {input_dir}, "
            f"pero se encontraron {len(files)}. "
            f"Formatos validos: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    return files


def read_all_inputs(files: List[Path]) -> str:
    blocks: List[str] = []
    for index, path in enumerate(files, start=1):
        try:
            text = extract_input_text(path)
        except Exception as exc:
            raise RuntimeError(
                f"Error leyendo el archivo {path.name}: {exc}"
            ) from exc
        if not text.strip():
            raise RuntimeError(
                f"El archivo {path.name} no tiene contenido textual extraible."
            )
        if len(text) > MAX_CONTENT_CHARS_PER_FILE:
            text = text[:MAX_CONTENT_CHARS_PER_FILE] + "\n[...contenido truncado...]"
        blocks.append(f"===ARCHIVO {index}: {path.name}===\n{text}")
    return "\n\n".join(blocks)


REQUIRED_STRUCTURE = """ESTRUCTURA OBLIGATORIA QUE DEBES SEGUIR EN CADA DOCUMENTO:

ACA (Proyecto Final / Espacio de Actividad Final), en este orden exacto:
1. TITULO DEL PROYECTO
2. INTRODUCCION Y CONTEXTUALIZACION
3. OBJETIVOS DE APRENDIZAJE  (verbos de accion; NO mencionar "Bloom")
4. DESCRIPCION DE LA ACTIVIDAD  (6 a 8 pasos numerados, detallados)
5. FORMATO DE ENTREGA  (PDF, Arial 12, interlineado 1.5, margenes 2.5 cm, 6-8 paginas, APA 7)
6. BIBLIOGRAFIA  (6-10 referencias APA 7 completas y reales)

PRESENTACION DE LA ASIGNATURA, en este orden exacto:
1. TEMAS DE LA ASIGNATURA  (5 temas en vinietas, uno por documento G1-G5)
2. GANCHO DE BIENVENIDA (HOOK)
3. PROPOSITO FORMATIVO
4. RESUMEN DE CONTENIDOS (EJES TEMATICOS)
   Subtitulos obligatorios en este orden:
   - Fundamentos
   - Factor Humano / Comunitario
   - Estandarizacion / Planificacion
   - Sistemas o Gestion especifica de la asignatura
5. VALOR AGREGADO PROFESIONAL
6. LLAMADOS A LA ACCION (CTA), exactamente estos tres en este orden:
   - Leer los documentos G1 a G5 antes de iniciar las actividades
   - Participar activamente en el foro de integracion
   - Completar el cuestionario diagnostico de ingreso en Moodle

FORO ACADEMICO, en este orden exacto:
1. CONTEXTUALIZACION DE LA TEMATICA
2. PREGUNTA INTEGRADORA  (una sola pregunta, abierta y critica)
3. INSTRUCCIONES DE PARTICIPACION
   - APORTE INICIAL  (200-300 palabras, argumentado y referenciado)
   - INTERACCION OBLIGATORIA  (debe decir literalmente "al menos 5 companeros")
   - CALIDAD DEL DEBATE
4. REQUERIMIENTOS DE FORMATO Y BIBLIOGRAFIA
   - Cierra con una seccion "REFERENCIAS" que liste TRES referencias APA 7 reales, completas y verificables."""


def build_user_prompt(
    combined_text: str,
    subject: str,
    program: str,
) -> str:
    return f"""Asignatura: {subject}
Programa: {program}

{REQUIRED_STRUCTURE}

A continuacion estan los 5 documentos fuente extraidos del usuario:

{combined_text}

Tu tarea: ejecuta el prompt maestro (los tres pasos: ACA, PRESENTACION, FORO) usando
toda la informacion anterior y la asignatura/programa indicados, respetando la
estructura obligatoria de cada documento.

Formato de respuesta OBLIGATORIO. Responde unicamente con tres bloques delimitados
asi, sin texto antes ni despues:

===ACA===
[contenido completo del Espacio de Actividad Final / Proyecto Final]

===PRESENTACION===
[contenido completo de la Presentacion de la Asignatura]

===FORO===
[contenido completo del Foro academico]

Reglas adicionales de salida para parseo correcto:
- Escribe los titulos principales y subtitulos en MAYUSCULAS, en una sola linea, sin numeracion decorativa.
- Usa el guion "-" como vinieta cuando corresponda.
- No uses asteriscos, markdown ni emojis.
- No incluyas comentarios sobre el proceso ni cierres con frases tipo "Espero que sea util".
- Manten un tono academico, claro y listo para entrega institucional.
""".strip()


def call_openai(
    client,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int,
    temperature: float,
) -> str:
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=temperature,
            max_tokens=max_tokens,
        )
    except Exception as exc:
        raise RuntimeError(f"OpenAI fallo al generar la respuesta: {exc}") from exc
    return (response.choices[0].message.content or "").strip()


def split_response(response: str) -> Dict[str, str]:
    pattern = re.compile(r"===\s*(ACA|PRESENTACION|FORO)\s*===", re.IGNORECASE)
    matches = list(pattern.finditer(response))
    if len(matches) < 3:
        found = [match.group(1).upper() for match in matches]
        raise ValueError(
            "La IA no devolvio los 3 bloques esperados "
            "(===ACA===, ===PRESENTACION===, ===FORO===). "
            f"Bloques detectados: {found}"
        )
    blocks: Dict[str, str] = {}
    for index, match in enumerate(matches):
        key = match.group(1).upper()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(response)
        blocks[key] = response[start:end].strip()
    for required in DOCUMENT_TYPES:
        if not blocks.get(required):
            raise ValueError(
                f"El bloque {required} esta vacio o no fue devuelto por la IA."
            )
    return blocks


REQUIRED_SECTIONS: Dict[str, Tuple[str, ...]] = {
    "ACA": (
        "TITULO DEL PROYECTO",
        "INTRODUCCION Y CONTEXTUALIZACION",
        "OBJETIVOS DE APRENDIZAJE",
        "DESCRIPCION DE LA ACTIVIDAD",
        "FORMATO DE ENTREGA",
        "BIBLIOGRAFIA",
    ),
    "PRESENTACION": (
        "TEMAS DE LA ASIGNATURA",
        "GANCHO DE BIENVENIDA",
        "PROPOSITO FORMATIVO",
        "RESUMEN DE CONTENIDOS",
        "VALOR AGREGADO PROFESIONAL",
        "LLAMADOS A LA ACCION",
    ),
    "FORO": (
        "CONTEXTUALIZACION DE LA TEMATICA",
        "PREGUNTA INTEGRADORA",
        "INSTRUCCIONES DE PARTICIPACION",
        "REQUERIMIENTOS DE FORMATO",
    ),
}

CRITICAL_WARNING_PATTERNS = (
    "Falta la seccion obligatoria",
    "Solo se detectaron 0",
    "no fue devuelto",
    "esta vacio",
)

ACA_CRITICAL_KEYWORDS = ("BIBLIOGRAFIA", "FORMATO DE ENTREGA")
FORO_CRITICAL_KEYWORDS = ("al menos 5 companeros", "CINCO COMPANEROS", "REFERENCIAS")
PRESENTACION_CRITICAL_KEYWORDS = ("Temas de la asignatura", "Resumen de Contenidos", "LLAMADOS A LA ACCION")

MAX_REPAIR_ATTEMPTS = 2


def _is_critical_warning(doc_type: str, warning: str) -> bool:
    normalized = warning.upper()
    if any(p.upper() in normalized for p in CRITICAL_WARNING_PATTERNS):
        return True
    if doc_type == "ACA" and any(k.upper() in normalized for k in ACA_CRITICAL_KEYWORDS):
        return True
    if doc_type == "FORO" and any(k.upper() in normalized for k in FORO_CRITICAL_KEYWORDS):
        return True
    if doc_type == "PRESENTACION" and any(k.upper() in normalized for k in PRESENTACION_CRITICAL_KEYWORDS):
        return True
    if "0 REFERENCIAS" in normalized or "0 TEMAS" in normalized or "0 POSIBLES" in normalized:
        return True
    return False


def classify_warnings(doc_type: str, warnings: List[str]) -> Tuple[List[str], List[str]]:
    critical = [w for w in warnings if _is_critical_warning(doc_type, w)]
    minor = [w for w in warnings if not _is_critical_warning(doc_type, w)]
    return critical, minor


def build_repair_prompt(
    doc_type: str,
    current_content: str,
    critical_issues: List[str],
    combined_text: str,
    subject: str,
    program: str,
) -> str:
    issues_text = "\n".join(f"- {issue}" for issue in critical_issues)
    return f"""REPARACION OBLIGATORIA - Documento: {doc_type}

El documento generado tiene los siguientes problemas CRITICOS que debes corregir:

{issues_text}

INSTRUCCIONES DE REPARACION:
1. Genera el documento {doc_type} COMPLETO desde cero.
2. NO omitas ninguna seccion obligatoria.
3. Cumple TODAS las longitudes minimas especificadas en el prompt original.
4. Incluye TODAS las keywords obligatorias.
5. Mantén el contenido academico y profesional del documento original.
6. Sigue EXACTAMENTE la estructura contractual del prompt original.

Asignatura: {subject}
Programa: {program}

Documentos fuente de referencia:
{combined_text[:15000]}

Genera el documento {doc_type} COMPLETO y CORREGIDO, sin texto adicional antes ni despues.
""".strip()

PRESENTATION_AXES: Tuple[str, ...] = (
    "FUNDAMENTOS",
    "FACTOR HUMANO",
    "ESTANDARIZACION",
    "SISTEMAS",
)

PRESENTATION_CTA_KEYWORDS: Tuple[str, ...] = (
    "G1",
    "FORO DE INTEGRACION",
    "CUESTIONARIO DIAGNOSTICO",
    "MOODLE",
)


def _normalize_for_check(value: str) -> str:
    return _strip_accents(value or "").upper()


def _section_block(content: str, section_marker: str, next_markers: Tuple[str, ...]) -> str:
    """Devuelve el contenido entre `section_marker` y el siguiente marcador conocido."""
    normalized = _normalize_for_check(content)
    start_match = re.search(re.escape(section_marker), normalized)
    if not start_match:
        return ""
    start = start_match.end()
    end = len(content)
    for marker in next_markers:
        next_match = re.search(re.escape(marker), normalized[start:])
        if next_match:
            candidate_end = start + next_match.start()
            if candidate_end < end:
                end = candidate_end
    return content[start:end]


def _count_apa_references(text: str) -> int:
    references = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if re.search(r"\((?:19|20)\d{2}[a-z]?\)", stripped) and len(stripped) > 30:
            references += 1
    return references


def _count_bullet_lines(text: str) -> int:
    bullet_chars = ("•", "●", "○", "·", "*", "-", "—", "–")
    count = 0
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(bullet_chars) and len(stripped) > 2:
            count += 1
    return count


def validate_blocks(blocks: Dict[str, str]) -> List[str]:
    """Devuelve una lista de advertencias detectadas. Lista vacia = todo OK.

    NOTE: This function is for the OLD sequential mode where all 3 docs
    were generated in a single response. For parallel mode, use
    validate_single_docx() instead.
    """
    warnings: List[str] = []

    for doc_type, sections in REQUIRED_SECTIONS.items():
        normalized = _normalize_for_check(blocks.get(doc_type, ""))
        for section in sections:
            if section not in normalized:
                warnings.append(f"[{doc_type}] Falta la seccion obligatoria: '{section}'")

    foro = blocks.get("FORO", "")
    foro_norm = _normalize_for_check(foro)
    if "AL MENOS 5 COMPANEROS" not in foro_norm and "CINCO COMPANEROS" not in foro_norm:
        warnings.append(
            "[FORO] No se encontro la regla 'al menos 5 companeros' / 'cinco companeros' en Interaccion Obligatoria."
        )

    apa_in_foro = _count_apa_references(foro)
    if apa_in_foro < 3:
        warnings.append(
            f"[FORO] Solo se detectaron {apa_in_foro} posibles referencias APA al final. Se requieren al menos 3."
        )

    presentation = blocks.get("PRESENTACION", "")
    temas_block = _section_block(
        presentation,
        "TEMAS DE LA ASIGNATURA",
        ("GANCHO DE BIENVENIDA", "PROPOSITO FORMATIVO"),
    )
    bullet_count = _count_bullet_lines(temas_block)
    if bullet_count < 5:
        warnings.append(
            f"[PRESENTACION] Solo se detectaron {bullet_count} temas listados en 'Temas de la asignatura'. Se esperan 5."
        )

    resumen_block = _section_block(
        presentation,
        "RESUMEN DE CONTENIDOS",
        ("VALOR AGREGADO PROFESIONAL", "LLAMADOS A LA ACCION"),
    )
    resumen_norm = _normalize_for_check(resumen_block)
    missing_axes = [axis for axis in PRESENTATION_AXES if axis not in resumen_norm]
    if missing_axes:
        warnings.append(
            "[PRESENTACION] Faltan ejes en 'Resumen de Contenidos': "
            + ", ".join(missing_axes)
        )

    cta_block = _section_block(
        presentation,
        "LLAMADOS A LA ACCION",
        (),
    )
    cta_norm = _normalize_for_check(cta_block)
    missing_cta = [keyword for keyword in PRESENTATION_CTA_KEYWORDS if keyword not in cta_norm]
    if missing_cta:
        warnings.append(
            "[PRESENTACION] Faltan referencias clave en CTA (G1, foro de integracion, cuestionario, Moodle): "
            + ", ".join(missing_cta)
        )

    aca = blocks.get("ACA", "")
    aca_norm = _normalize_for_check(aca)
    if "BLOOM" in aca_norm:
        warnings.append("[ACA] Aparece la palabra 'Bloom' en el documento; debe omitirse segun la regla.")

    formato_block = _section_block(
        aca,
        "FORMATO DE ENTREGA",
        ("BIBLIOGRAFIA",),
    )
    formato_norm = _normalize_for_check(formato_block)
    expected_formato_tokens = ("PDF", "ARIAL", "1.5", "APA")
    missing_formato = [token for token in expected_formato_tokens if token not in formato_norm]
    if missing_formato:
        warnings.append(
            "[ACA] 'Formato de entrega' no incluye todos los requisitos (PDF, Arial, 1.5, APA). Falta: "
            + ", ".join(missing_formato)
        )

    bibliografia_block = _section_block(aca, "BIBLIOGRAFIA", ())
    bib_count = _count_apa_references(bibliografia_block)
    if bib_count < 6:
        warnings.append(
            f"[ACA] Solo se detectaron {bib_count} referencias en Bibliografia. Se esperan al menos 6."
        )

    return warnings


def validate_single_docx(doc_type: str, content: str) -> List[str]:
    """Validate a single DOCX document (for parallel mode).

    Only checks sections and rules relevant to the specific document type.
    """
    warnings: List[str] = []
    normalized = _normalize_for_check(content)

    # Check required sections for this doc type
    sections = REQUIRED_SECTIONS.get(doc_type, ())
    for section in sections:
        if section not in normalized:
            warnings.append(f"[{doc_type}] Falta la seccion obligatoria: '{section}'")

    if doc_type == "ACA":
        if "BLOOM" in normalized:
            warnings.append("[ACA] Aparece la palabra 'Bloom' en el documento; debe omitirse segun la regla.")

        formato_block = _section_block(content, "FORMATO DE ENTREGA", ("BIBLIOGRAFIA",))
        formato_norm = _normalize_for_check(formato_block)
        expected_formato_tokens = ("PDF", "ARIAL", "1.5", "APA")
        missing_formato = [token for token in expected_formato_tokens if token not in formato_norm]
        if missing_formato:
            warnings.append(
                "[ACA] 'Formato de entrega' no incluye todos los requisitos (PDF, Arial, 1.5, APA). Falta: "
                + ", ".join(missing_formato)
            )

        bibliografia_block = _section_block(content, "BIBLIOGRAFIA", ())
        bib_count = _count_apa_references(bibliografia_block)
        if bib_count < 6:
            warnings.append(
                f"[ACA] Solo se detectaron {bib_count} referencias en Bibliografia. Se esperan al menos 6."
            )

    elif doc_type == "FORO":
        if "AL MENOS 5 COMPANEROS" not in normalized and "CINCO COMPANEROS" not in normalized:
            warnings.append(
                "[FORO] No se encontro la regla 'al menos 5 companeros' / 'cinco companeros' en Interaccion Obligatoria."
            )

        apa_count = _count_apa_references(content)
        if apa_count < 3:
            warnings.append(
                f"[FORO] Solo se detectaron {apa_count} posibles referencias APA al final. Se requieren al menos 3."
            )

    elif doc_type == "PRESENTACION":
        temas_block = _section_block(
            content,
            "TEMAS DE LA ASIGNATURA",
            ("GANCHO DE BIENVENIDA", "PROPOSITO FORMATIVO"),
        )
        bullet_count = _count_bullet_lines(temas_block)
        if bullet_count < 5:
            warnings.append(
                f"[PRESENTACION] Solo se detectaron {bullet_count} temas listados en 'Temas de la asignatura'. Se esperan 5."
            )

        resumen_block = _section_block(
            content,
            "RESUMEN DE CONTENIDOS",
            ("VALOR AGREGADO PROFESIONAL", "LLAMADOS A LA ACCION"),
        )
        resumen_norm = _normalize_for_check(resumen_block)
        missing_axes = [axis for axis in PRESENTATION_AXES if axis not in resumen_norm]
        if missing_axes:
            warnings.append(
                "[PRESENTACION] Faltan ejes en 'Resumen de Contenidos': "
                + ", ".join(missing_axes)
            )

        cta_block = _section_block(
            content,
            "LLAMADOS A LA ACCION",
            (),
        )
        cta_norm = _normalize_for_check(cta_block)
        missing_cta = [keyword for keyword in PRESENTATION_CTA_KEYWORDS if keyword not in cta_norm]
        if missing_cta:
            warnings.append(
                "[PRESENTACION] Faltan referencias clave en CTA (G1, foro de integracion, cuestionario, Moodle): "
                + ", ".join(missing_cta)
            )

    return warnings


def normalize_filename_part(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", "", value or "").strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned.upper() or "GENERICO"


def build_output_filename(doc_type: str, subject: str, program: str) -> str:
    subject_part = normalize_filename_part(subject)
    program_part = normalize_filename_part(program)
    return f"{doc_type}_{subject_part}_{program_part}.docx"


def _force_run_font(run, *, bold: bool, size_pt: int) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _configure_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    _force_run_font(run, bold=True, size_pt=HEADING_FONT_SIZE)


def _add_paragraph(doc: Document, text: str, *, bullet: bool = False) -> None:
    if bullet:
        paragraph = doc.add_paragraph(style="List Bullet")
    else:
        paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    _force_run_font(run, bold=False, size_pt=BODY_FONT_SIZE)


def _strip_accents(value: str) -> str:
    nfkd = unicodedata.normalize("NFKD", value)
    return "".join(ch for ch in nfkd if not unicodedata.combining(ch))


def _is_heading_line(text: str) -> bool:
    stripped = text.strip(" :")
    if not stripped or len(stripped) > 110:
        return False
    ascii_form = _strip_accents(stripped)
    letters = [ch for ch in ascii_form if ch.isalpha()]
    if not letters:
        return False
    if all(ch.isupper() for ch in letters) and len(letters) >= 3:
        return True
    if re.match(r"^\d+[\.\)]\s+[A-ZÁÉÍÓÚÑ]", stripped) and len(stripped.split()) <= 14:
        return True
    if text.strip().endswith(":") and len(stripped.split()) <= 12:
        first_letter = next((ch for ch in stripped if ch.isalpha()), "")
        if first_letter.isupper():
            return True
    return False


def _detect_bullet(text: str) -> Tuple[bool, str]:
    stripped = text.strip()
    if not stripped:
        return False, ""
    bullet_chars = ("•", "●", "○", "·", "*", "-", "—", "–")
    if stripped.startswith(bullet_chars):
        return True, stripped.lstrip("".join(bullet_chars) + " \t").strip()
    return False, stripped


def render_docx(content: str, output_path: Path, document_title: str) -> None:
    doc = Document()
    _configure_default_styles(doc)

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(6)
    title_run = title_paragraph.add_run(document_title)
    _force_run_font(title_run, bold=True, size_pt=HEADING_FONT_SIZE)

    for raw_block in re.split(r"\n\s*\n", content.strip()):
        block = raw_block.strip()
        if not block:
            continue
        for raw_line in block.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue
            is_bullet, body_text = _detect_bullet(line)
            if is_bullet:
                _add_paragraph(doc, body_text, bullet=True)
            elif _is_heading_line(line):
                _add_heading(doc, line.strip(" :"))
            else:
                _add_paragraph(doc, line.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def load_docx_prompt(doc_type: str, base_dir: Optional[Path] = None) -> str:
    try:
        prompt_path = resolve_docx_prompt(doc_type, base_dir)
        return prompt_path.read_text(encoding="utf-8")
    except FileNotFoundError:
        # Fallback chain for legacy paths
        fallback_candidates = [
            PROJECT_ROOT / "prompts" / "docx" / "system_prompt_documentos_academicos.md",
            PROJECT_ROOT / "prompts" / "system_prompt_documentos_academicos.md",
        ]
        for candidate in fallback_candidates:
            if candidate.exists():
                return candidate.read_text(encoding="utf-8")
        raise FileNotFoundError(
            f"No se encontro ningun prompt DOCX para '{doc_type}'. "
            f"Buscado en: {[str(c) for c in fallback_candidates]}"
        )


def build_independent_docx_prompt(
    doc_type: str,
    combined_text: str,
    subject: str,
    program: str,
) -> Tuple[str, str]:
    system_prompt = load_docx_prompt(doc_type)
    user_prompt = f"""Asignatura: {subject}
Programa: {program}

A continuacion estan los documentos fuente extraidos:

{combined_text}

Genera unicamente el contenido del documento {doc_type} solicitado, sin texto adicional antes ni despues.
""".strip()
    return system_prompt, user_prompt


def generate_single_docx(
    client,
    model: str,
    doc_type: str,
    combined_text: str,
    subject: str,
    program: str,
    max_tokens: int,
    temperature: float,
    output_dir: Path,
    max_repair_attempts: int = MAX_REPAIR_ATTEMPTS,
) -> Dict[str, object]:
    start_time = time.time()
    result = {
        "task": doc_type,
        "status": "pending",
        "output_file": "",
        "warnings": [],
        "warnings_critical": [],
        "warnings_minor": [],
        "duration_seconds": 0,
        "word_count": 0,
        "error": "",
        "repair_attempts": 0,
        "repaired_successfully": False,
        "repair_duration_seconds": 0,
    }
    repair_start = 0.0
    try:
        system_prompt, user_prompt = build_independent_docx_prompt(doc_type, combined_text, subject, program)
        response_text = call_openai(
            client=client,
            model=model,
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        all_warnings = validate_single_docx(doc_type, response_text)
        critical_issues, minor_issues = classify_warnings(doc_type, all_warnings)

        if critical_issues and max_repair_attempts > 0:
            print(f"  [{doc_type}] validation failed - {len(critical_issues)} critical issues detected")
            repair_start = time.time()
            for attempt in range(1, max_repair_attempts + 1):
                print(f"  [{doc_type}] repair attempt {attempt}/{max_repair_attempts}")
                repair_prompt = build_repair_prompt(
                    doc_type=doc_type,
                    current_content=response_text,
                    critical_issues=critical_issues,
                    combined_text=combined_text,
                    subject=subject,
                    program=program,
                )
                repaired_text = call_openai(
                    client=client,
                    model=model,
                    system_prompt=system_prompt,
                    user_prompt=repair_prompt,
                    max_tokens=max_tokens,
                    temperature=min(temperature, 0.3),
                )
                repaired_warnings = validate_single_docx(doc_type, repaired_text)
                repaired_critical, repaired_minor = classify_warnings(doc_type, repaired_warnings)
                if not repaired_critical:
                    response_text = repaired_text
                    all_warnings = repaired_warnings
                    critical_issues = []
                    minor_issues = repaired_minor
                    result["repaired_successfully"] = True
                    print(f"  [{doc_type}] repaired successfully on attempt {attempt}")
                    break
                else:
                    print(f"  [{doc_type}] repair attempt {attempt} still has {len(repaired_critical)} critical issues")
                    critical_issues = repaired_critical
                    minor_issues = repaired_minor
            result["repair_attempts"] = max_repair_attempts if critical_issues else sum(1 for _ in range(max_repair_attempts))
            result["repair_duration_seconds"] = round(time.time() - repair_start, 2)

        filename = build_output_filename(doc_type, subject, program)
        output_path = output_dir / filename
        title = f"{DOCUMENT_TITLES[doc_type]} - {subject.upper()}"
        render_docx(response_text, output_path, title)
        result["status"] = "success"
        result["output_file"] = filename
        result["word_count"] = word_count(response_text)
        result["warnings"] = all_warnings
        result["warnings_critical"] = critical_issues
        result["warnings_minor"] = minor_issues
    except Exception as exc:
        result["status"] = "error"
        result["error"] = str(exc)
    finally:
        result["duration_seconds"] = round(time.time() - start_time, 2)
    return result


def _first_field_value(raw: str) -> str:
    """Recorta el valor capturado al primer campo adicional (ej. '. PROGRAMA:')."""
    raw = clean_text(raw)
    cutoff = re.search(r"\.\s+[A-ZÁÉÍÓÚÑ]{2,}[^a-z]*:", raw)
    if cutoff:
        raw = raw[: cutoff.start()].strip(" .")
    return raw


def infer_subject_and_program(combined_text: str) -> tuple:
    """Busca en el texto combinado etiquetas comunes de sílabos y guiones."""
    subject = DEFAULT_SUBJECT
    program = DEFAULT_PROGRAM

    # Encabezado propio de guiones generados por generate_guiones.py:
    # "CONTENIDO: tema. ASIGNATURA: xxx. PROGRAMA: yyy. ..."
    guion_header = re.search(
        r"ASIGNATURA\s*:\s*([^.]{3,80})\.\s*PROGRAMA\s*:\s*([^.]{3,80})",
        combined_text,
        re.IGNORECASE,
    )
    if guion_header:
        subject = clean_text(guion_header.group(1))
        program = clean_text(guion_header.group(2))
        return subject, program

    subject_pattern = re.compile(
        r"(?:asignatura|nombre de la asignatura|materia|curso)\s*[:\-]?\s*([^\n\r]{3,120})",
        re.IGNORECASE,
    )
    program_pattern = re.compile(
        r"(?:programa|carrera|plan de estudios)\s*[:\-]?\s*([^\n\r]{3,120})",
        re.IGNORECASE,
    )

    subject_match = subject_pattern.search(combined_text)
    if subject_match:
        candidate = _first_field_value(subject_match.group(1))
        if 3 <= len(candidate) <= 80:
            subject = candidate

    program_match = program_pattern.search(combined_text)
    if program_match:
        candidate = _first_field_value(program_match.group(1))
        if 3 <= len(candidate) <= 80:
            program = candidate

    return subject, program


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Genera 3 documentos academicos (ACA, PRESENTACION, FORO) en .docx "
            "a partir de exactamente 5 archivos fuente analizados con OpenAI."
        )
    )
    parser.add_argument(
        "--input-dir",
        default=str(DEFAULT_INPUT_DIR),
        help=(
            f"Carpeta con exactamente 5 archivos fuente (.pdf, .docx o .txt). "
            f"Por defecto: {DEFAULT_INPUT_DIR}"
        ),
    )
    parser.add_argument(
        "--output-dir",
        default=str(DEFAULT_OUTPUT_DIR),
        help=(
            f"Carpeta donde se guardaran los .docx generados. "
            f"Por defecto: {DEFAULT_OUTPUT_DIR}"
        ),
    )
    parser.add_argument(
        "--subject",
        default="",
        help=(
            "Nombre de la asignatura. Si se omite, se intenta inferir del "
            f"contenido de los archivos o se usa '{DEFAULT_SUBJECT}'."
        ),
    )
    parser.add_argument(
        "--program",
        default="",
        help=(
            "Nombre del programa academico. Si se omite, se intenta inferir del "
            f"contenido de los archivos o se usa '{DEFAULT_PROGRAM}'."
        ),
    )
    parser.add_argument(
        "--prompt",
        default=str(DEFAULT_PROMPT_PATH),
        help="Ruta al prompt maestro (.md)",
    )
    parser.add_argument(
        "--model",
        default=os.getenv("OPENAI_MODEL", "gpt-4o"),
        help="Modelo OpenAI a utilizar",
    )
    parser.add_argument(
        "--max-tokens",
        type=int,
        default=6000,
        help="Maximo de tokens en la respuesta del modelo",
    )
    parser.add_argument(
        "--temperature",
        type=float,
        default=0.6,
        help="Creatividad de generacion",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Lee los archivos y construye los prompts sin llamar a la API",
    )
    return parser.parse_args()


def main() -> None:
    if load_dotenv:
        load_dotenv()

    args = parse_args()
    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)
    prompt_path = Path(args.prompt)

    if not prompt_path.exists():
        raise FileNotFoundError(
            f"No se encontro el prompt maestro: {prompt_path}. "
            "Asegurate de tener prompts/system_prompt_documentos_academicos.md."
        )

    files = collect_input_files(input_dir)
    print("Archivos detectados:")
    for path in files:
        print(f"  - {path.name}")

    combined_text = read_all_inputs(files)

    subject = args.subject.strip()
    program = args.program.strip()
    if not subject or not program:
        inferred_subject, inferred_program = infer_subject_and_program(combined_text)
        if not subject:
            subject = inferred_subject
            print(f"Asignatura inferida: {subject}")
        if not program:
            program = inferred_program
            print(f"Programa inferido: {program}")

    system_prompt = load_system_prompt(prompt_path)
    user_prompt = build_user_prompt(
        combined_text=combined_text,
        subject=subject,
        program=program,
    )

    output_dir.mkdir(parents=True, exist_ok=True)

    if args.dry_run:
        print("\nDry-run activo. No se llamo a la API.")
        print(f"Caracteres totales del contenido combinado: {len(combined_text)}")
        print(f"Caracteres del user prompt: {len(user_prompt)}")
        return

    if OpenAI is None:
        raise RuntimeError(
            "Falta instalar el paquete openai. Ejecuta: pip install -r requirements.txt"
        )
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("Falta OPENAI_API_KEY en variables de entorno.")

    client = OpenAI()
    print(f"\nLlamando al modelo {args.model}...")
    response = call_openai(
        client=client,
        model=args.model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        max_tokens=args.max_tokens,
        temperature=args.temperature,
    )

    blocks = split_response(response)

    for doc_type in DOCUMENT_TYPES:
        filename = build_output_filename(doc_type, subject, program)
        output_path = output_dir / filename
        title = f"{DOCUMENT_TITLES[doc_type]} - {subject.upper()}"
        render_docx(blocks[doc_type], output_path, title)
        print(f"Guardado: {output_path}")

    warnings = validate_blocks(blocks)
    if warnings:
        print("\n=== ADVERTENCIAS DE VALIDACION ===")
        for warning in warnings:
            print(f"  ! {warning}")
        print("=== FIN DE ADVERTENCIAS ===")
        print(
            "Los archivos se generaron pero algunas reglas criticas del prompt no se cumplen al 100%. "
            "Revisalos manualmente o vuelve a ejecutar."
        )
    else:
        print("\nValidacion OK: los 3 documentos cumplen los requisitos minimos.")

    print("\nGeneracion completa.")


if __name__ == "__main__":
    try:
        main()
    except (FileNotFoundError, ValueError, RuntimeError) as error:
        print(f"ERROR: {error}", file=sys.stderr)
        sys.exit(1)
