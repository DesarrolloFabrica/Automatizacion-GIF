from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

ENGINE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = ENGINE_DIR.parent
sys.path.insert(0, str(PROJECT_ROOT))

from automation_engine.generate_materiales_categoria import generate_all_materiales
from automation_engine.utils.naming import normalize_for_filename


DEFAULT_INPUT_DIR = PROJECT_ROOT / "drive-download-20260511T175951Z-3-001"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "outputs" / "material03_preview"


def detect_granule_code(path: Path) -> str | None:
    match = re.search(r"(?:^|[^A-Z0-9])G\s*([1-5])(?=$|[^0-9])", path.stem.upper())
    if not match:
        return None
    return f"G{match.group(1)}"


def topic_from_filename(path: Path, code: str) -> str:
    stem = re.sub(rf"(?i)(?:^|[^A-Z0-9])G\s*{code[1]}(?=$|[^0-9])", " ", path.stem, count=1)
    stem = re.sub(r"[_\-]+", " ", stem).strip()
    return normalize_for_filename(stem or code)


def discover_manual_granules(input_dir: Path) -> list[dict]:
    if not input_dir.exists() or not input_dir.is_dir():
        raise FileNotFoundError(f"La carpeta input no existe: {input_dir}")

    by_code: dict[str, Path] = {}
    for path in input_dir.rglob("*.docx"):
        if path.name.startswith("~$"):
            continue
        code = detect_granule_code(path)
        if not code:
            continue
        if code in by_code:
            raise ValueError(f"Granulo duplicado para {code}: {by_code[code]} y {path}")
        by_code[code] = path

    expected = {f"G{i}" for i in range(1, 6)}
    found = set(by_code)
    missing = sorted(expected - found)
    extra = sorted(found - expected)
    if missing or extra:
        details = []
        if missing:
            details.append(f"faltan: {', '.join(missing)}")
        if extra:
            details.append(f"sobran: {', '.join(extra)}")
        raise ValueError(f"Se esperaban exactamente 5 granulos G1-G5 ({'; '.join(details)}).")

    return [
        {"code": code, "path": by_code[code], "topic": topic_from_filename(by_code[code], code)}
        for code in sorted(by_code, key=lambda value: int(value[1:]))
    ]


def prepare_staging(granules: list[dict], output_dir: Path) -> Path:
    staging_dir = output_dir / "_staging_granules"
    if staging_dir.exists():
        shutil.rmtree(staging_dir)
    staging_dir.mkdir(parents=True, exist_ok=True)

    for granule in granules:
        target = staging_dir / f"{granule['code']}_{granule['topic']}.docx"
        shutil.copy2(granule["path"], target)
        granule["staged_path"] = target
    return staging_dir


def clean_previous_preview_files(output_dir: Path) -> None:
    for path in output_dir.glob("03_G*_VIDEO_PRESENTACION_DEL_PROBLEMA_*.docx"):
        path.unlink()
    for path in output_dir.glob("debug_*_G*_03.*"):
        path.unlink()


def collect_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return value.lower()


def validate_material03_output(output_dir: Path) -> list[dict]:
    files = sorted(output_dir.glob("03_G*_VIDEO_PRESENTACION_DEL_PROBLEMA_*.docx"))
    validations = []
    if len(files) != 5:
        raise ValueError(f"Se esperaban 5 DOCX de material 03 y se encontraron {len(files)} en {output_dir}.")

    for path in files:
        doc = Document(path)
        text = collect_docx_text(path)
        scene_count = len(re.findall(r"\bEscena\s+\d+\b", text, flags=re.IGNORECASE))
        normalized_text = normalize_text(text)
        has_scene_table = all(term in normalized_text for term in ["locucion", "texto en pantalla", "transicion"])
        body_chars = len(text)
        validation = {
            "file": path.name,
            "size_bytes": path.stat().st_size,
            "body_chars": body_chars,
            "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
            "tables": len(doc.tables),
            "scene_count": scene_count,
            "estimated_multiple_pages": body_chars >= 4500 and scene_count >= 8,
            "not_cover_only": body_chars >= 2500 and scene_count >= 8,
            "has_complete_scene_structure": scene_count >= 8 and has_scene_table,
        }
        validations.append(validation)
        if not validation["not_cover_only"]:
            raise ValueError(f"DOCX posiblemente vacio o solo portada: {path.name}")
        if not validation["has_complete_scene_structure"]:
            raise ValueError(f"DOCX sin escenas completas suficientes: {path.name}")

    return validations


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Flujo local de prueba: genera solo material 03 de Pregrado desde granulos DOCX existentes."
    )
    parser.add_argument("--input-dir", default=str(DEFAULT_INPUT_DIR), help="Carpeta raiz con los DOCX G1-G5 descargados manualmente")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR), help="Carpeta local donde se guardan solo los DOCX del material 03")
    parser.add_argument("--model", default=os.getenv("OPENAI_MODEL", "gpt-4o"), help="Modelo OpenAI")
    parser.add_argument("--max-tokens", type=int, default=8000, help="Maximo de tokens por material")
    parser.add_argument("--temperature", type=float, default=0.5, help="Creatividad de generacion")
    return parser.parse_args()


def main() -> None:
    if load_dotenv:
        load_dotenv()

    args = parse_args()
    input_dir = Path(args.input_dir).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    os.environ.pop("AUTOMATIZACION_GIF_JOB_ID", None)

    print("=== PREVIEW LOCAL MATERIAL 03 - PREGRADO ===")
    print(f"Input: {input_dir}")
    print(f"Output: {output_dir}")
    print("Fases omitidas: generate_guiones, pipeline_local, ZIP, Drive sync, actividades Moodle y otros materiales.")

    granules = discover_manual_granules(input_dir)
    print(f"Granulos detectados: {len(granules)}")
    for granule in granules:
        print(f"  - {granule['code']}: {granule['path'].name}")

    clean_previous_preview_files(output_dir)
    staging_dir = prepare_staging(granules, output_dir)
    print(f"Staging local: {staging_dir}")

    result = generate_all_materiales(
        job_id="material03_preview",
        category_key="pregrado",
        generated_dir=staging_dir,
        output_base=output_dir,
        model=args.model,
        max_tokens=args.max_tokens,
        temperature=args.temperature,
        only_materials={"03"},
        flat_output=True,
        disable_drive_upload=True,
        debug_dir=output_dir,
    )

    validations = validate_material03_output(output_dir)
    preview_summary = {
        "job_id": "material03_preview",
        "fecha": datetime.now(timezone.utc).isoformat(),
        "input_dir": str(input_dir),
        "output_dir": str(output_dir),
        "material_forzado": "03",
        "granules_detected": [{"code": g["code"], "source": str(g["path"])} for g in granules],
        "generation_summary": result["summary"],
        "validations": validations,
    }
    summary_path = output_dir / "preview_summary.json"
    summary_path.write_text(json.dumps(preview_summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Validacion guardada: {summary_path}")
    print("Preview local completado: solo material 03 generado.")


if __name__ == "__main__":
    main()
