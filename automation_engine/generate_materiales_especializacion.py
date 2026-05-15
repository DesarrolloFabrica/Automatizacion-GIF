from __future__ import annotations

import argparse
import json
import os
import re
import sys
import traceback
import unicodedata
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

ENGINE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = ENGINE_DIR.parent

sys.path.insert(0, str(PROJECT_ROOT))

from automation_engine.config.materiales_especializacion import (
    ESPECIALIZACION_PROMPT_PATH,
    MATERIALES_A_GENERAR,
    MATERIALES_RESERVADOS_FUTURO,
    VERSION_DEFECTO,
    EspecializacionConfig,
    MaterialConfig,
)
from automation_engine.utils.naming import (
    build_granule_folder_name,
    normalize_for_filename,
)
from automation_engine.utils.openai_client import get_openai_client, get_openai_model


EXPECTED_GRANULE_COUNT = 5
EXPECTED_MATERIAL_COUNT = 6
MIN_RESPONSE_CHARS = 200

FONT_NAME = "Arial"

COLOR_NAVY = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_BLUE = RGBColor(0x2E, 0x75, 0xB6)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
COLOR_MED_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B3A5C"
HEX_LIGHT_BLUE = "D6E4F0"
HEX_VERY_LIGHT = "F2F7FB"
HEX_ACCENT = "E8EEF4"
HEX_SEPARATOR = "B4C6E7"
HEX_WHITE = "FFFFFF"

COVER_TITLE_SIZE = Pt(24)
COVER_SUBTITLE_SIZE = Pt(14)
COVER_META_SIZE = Pt(11)
H1_SIZE = Pt(16)
H2_SIZE = Pt(13)
H3_SIZE = Pt(12)
BODY_SIZE = Pt(11)
BODY_SMALL = Pt(10)
FOOTER_SIZE = Pt(8)
TABLE_HEADER_SIZE = Pt(9)
TABLE_BODY_SIZE = Pt(10)

MATERIAL_VALIDATION_RULES = {
    "02": {"name": "FICHAS_DE_ESTUDIO_DE_EVIDENCIA", "type": "fichas", "min_fichas": 5},
    "03": {"name": "GLOSARIO_ESPECIALIZADO", "type": "glosario", "min_terminos": 12},
    "04": {"name": "REVISTA_DOSSIER", "type": "revista", "min_bloques": 14},
    "05": {"name": "INFOGRAFIA_MODELO_O_RUTA", "type": "infografia", "min_bloques": 7},
    "06": {"name": "PODCAST_DEBATE_EXPERTO", "type": "podcast", "min_segmentos": 9},
    "07": {"name": "VIDEO_SOLUCION_O_PROCEDIMIENTO", "type": "video", "min_escenas": 7},
}

MINIMALS_BY_CATEGORY: dict[str, dict[str, dict]] = {
    "especializacion": {
        "02": {"type": "fichas", "min_fichas": 5},
        "03": {"type": "glosario", "min_terminos": 12},
        "04": {"type": "revista", "min_bloques": 14},
        "05": {"type": "infografia", "min_bloques": 7},
        "06": {"type": "podcast", "min_segmentos": 9},
        "07": {"type": "video", "min_escenas": 7},
    },
    "pregrado": {
        "01": {"type": "podcast", "min_segmentos": 7},
        "02": {"type": "infografia", "min_bloques": 6},
        "03": {"type": "video", "min_escenas": 8},
        "04": {"type": "glosario", "min_terminos": 8},
        "05": {"type": "video", "min_escenas": 6},
        "06": {"type": "revista", "min_bloques": 12},
        "07": {"type": "fichas", "min_fichas": 3},
    },
}


def _get_validation_minimals(
    category_key: str | None,
    layout_nn: str | None,
    material_nn: str | None = None,
) -> dict:
    """Return validation minimums for a given category and material/layout number.

    Falls back to MATERIAL_VALIDATION_RULES defaults if category is unknown.
    """
    ck = (category_key or "").strip().lower()
    nn = (material_nn or layout_nn or "").strip()
    cat_minimals = MINIMALS_BY_CATEGORY.get(ck)
    if cat_minimals and nn in cat_minimals:
        return cat_minimals[nn]
    layout_key = (layout_nn or "").strip()
    return MATERIAL_VALIDATION_RULES.get(layout_key, {})

ARTIFACT_PATTERNS = [
    re.compile(r"^Datos recibidos\..*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^Generare unicamente el material solicitado\..*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^No generes los demas materiales\..*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^No agregues informacion que no este en el GUION MAESTRO\..*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^Entrega el contenido en formato tabla.*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^Confirmacion:.*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^A continuacion el material.*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^```(?:text|txt|markdown)?\s*$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^```\s*$", re.MULTILINE),
]


def _set_cell_shading(cell, color_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _set_cell_font(cell, size: Pt, bold: bool = False, color: RGBColor = COLOR_DARK_GRAY) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = _clean_visible_text(run.text)
            run.font.name = FONT_NAME
            run.font.size = size
            run.font.color.rgb = color
            run.bold = bold
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rfonts.set(qn(attr), FONT_NAME)


def _set_cell_borders(cell, top=None, bottom=None, start=None, end=None) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        if val:
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), val.get("val", "single"))
            elem.set(qn("w:sz"), val.get("sz", "4"))
            elem.set(qn("w:space"), val.get("space", "0"))
            elem.set(qn("w:color"), val.get("color", "B4C6E7"))
            tcBorders.append(elem)
    tcPr.append(tcBorders)


def _clean_visible_text(text: str) -> str:
    if not text:
        return ""
    cleaned = text.strip()
    cleaned = re.sub(r"<br\s*/?>", "\n", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bchecklist\b", "lista de criterios", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"</?(?:ul|ol)>", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"</?li>", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    return cleaned


def _add_text_block(doc, text, size=BODY_SIZE, color=COLOR_DARK_GRAY,
                    space_before=0, space_after=4):
    if not text:
        return
    raw = text.strip()
    li_matches = re.findall(r"<li[^>]*>(.*?)</li>", raw, flags=re.IGNORECASE | re.DOTALL)
    if li_matches:
        before_list = re.split(r"<ul[^>]*>|<ol[^>]*>", raw, maxsplit=1, flags=re.IGNORECASE)[0]
        before_list = _clean_visible_text(before_list)
        if before_list:
            _add_styled_para(doc, before_list, size=size, color=color,
                              space_before=space_before, space_after=space_after)
        for item in li_matches:
            _add_bullet_para(doc, _clean_visible_text(item), size=size, color=color)
        return

    cleaned = _clean_visible_text(raw)
    for para_text in cleaned.split("\n"):
        if para_text.strip():
            _add_styled_para(doc, para_text.strip(), size=size, color=color,
                              space_before=space_before, space_after=space_after)


def _add_bullet_para(doc, text, size=BODY_SIZE, color=COLOR_DARK_GRAY):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("\u2022  " + _clean_visible_text(text))
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.color.rgb = color


def _set_paragraph_spacing(para, before=0, after=6, line_spacing=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_border_below(para, color_hex="B4C6E7", sz="6"):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_styled_para(doc, text, size=BODY_SIZE, bold=False, color=COLOR_DARK_GRAY,
                      alignment=None, space_before=0, space_after=6, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(_clean_visible_text(text))
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if alignment:
        para.alignment = alignment
    _set_paragraph_spacing(para, before=space_before, after=space_after)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    return para


def _add_separator(doc):
    para = doc.add_paragraph()
    _add_border_below(para, color_hex=HEX_SEPARATOR, sz="4")
    _set_paragraph_spacing(para, before=8, after=8)


def _add_callout_box(doc, title, text):
    _add_styled_para(doc, "", space_before=8, space_after=0)
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    cell = box.rows[0].cells[0]
    cell.text = ""
    _set_cell_shading(cell, HEX_ACCENT)
    _set_cell_borders(cell,
        top={"val": "single", "sz": "6", "color": HEX_SEPARATOR},
        bottom={"val": "single", "sz": "6", "color": HEX_SEPARATOR},
        start={"val": "single", "sz": "6", "color": HEX_SEPARATOR},
        end={"val": "single", "sz": "6", "color": HEX_SEPARATOR},
    )
    if title:
        p = cell.paragraphs[0]
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = BODY_SMALL
        run.font.color.rgb = COLOR_NAVY
        _set_paragraph_spacing(p, before=6, after=4)
    if text:
        p2 = cell.add_paragraph()
        run2 = p2.add_run(_clean_visible_text(text))
        run2.font.name = FONT_NAME
        run2.font.size = BODY_SMALL
        run2.font.color.rgb = COLOR_DARK_GRAY
        _set_paragraph_spacing(p2, before=2, after=6)
    for width in [Cm(15)]:
        box.columns[0].width = width
    _add_styled_para(doc, "", space_before=4, space_after=0)


def _parse_markdown_tables(content: str) -> List[Tuple[List[str], List[List[str]]]]:
    tables = []
    current_header = None
    current_rows = []

    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if not cells:
                continue
            if _looks_like_table_header(cells):
                if current_header:
                    tables.append((current_header, current_rows))
                current_header = cells
                current_rows = []
            elif current_header:
                current_rows.append(cells)

    if current_header:
        tables.append((current_header, current_rows))

    return tables


def _looks_like_table_header(cells: List[str]) -> bool:
    normalized = " | ".join(_normalize_header_cell(c) for c in cells)
    header_signatures = [
        "campo | informacion",
        "ficha | titulo visible lado a",
        "no. | termino",
        "bloque | titulo visible | carga textual",
        "bloque | titulo visible | texto visible",
        "segmento | duracion estimada",
        "escena | duracion | funcion",
        "escena | duracion estimada | funcion",
        "escena | tiempo estimado | objetivo narrativo",
        "criterio | cumple | observacion",
        "referencia | uso en revista",
        "recurso sugerido | uso pedagogico",
        "elemento de edicion | indicacion",
        "recurso de accesibilidad | indicacion",
        "nota para diseno | contenido",
        "cierre integrado de la ruta | texto",
    ]
    return any(sig in normalized for sig in header_signatures)


def count_visible_docx_chars(doc: Document) -> int:
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return len("\n".join(parts))


def _add_raw_content_fallback(doc: Document, content: str) -> None:
    cleaned = clean_ai_response(content)
    if not cleaned:
        return
    _add_styled_para(doc, "Contenido generado", size=H1_SIZE, bold=True,
                      color=COLOR_NAVY, space_before=12, space_after=8)
    for line in cleaned.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            _add_styled_para(doc, _clean_visible_text(stripped), size=BODY_SMALL,
                              color=COLOR_DARK_GRAY, space_before=0, space_after=3)
        elif re.match(r"^#{1,6}\s+", stripped):
            _add_styled_para(doc, re.sub(r"^#{1,6}\s+", "", stripped), size=H2_SIZE,
                              bold=True, color=COLOR_NAVY, space_before=10, space_after=4)
        else:
            _add_text_block(doc, stripped, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=4)


def _normalize_header_cell(value: str) -> str:
    value = value.lower().strip()
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"\s+", " ", value)


def _header_map(header: List[str]) -> Dict[str, int]:
    return {_normalize_header_cell(h): i for i, h in enumerate(header)}


def _row_value(row: List[str], col_map: Dict[str, int], keys: List[str], fallback_idx: int = -1) -> str:
    for key in keys:
        idx = col_map.get(_normalize_header_cell(key))
        if idx is not None and idx < len(row):
            return row[idx]
    if 0 <= fallback_idx < len(row):
        return row[fallback_idx]
    return ""


def _split_table_header(lines: List[str]) -> Tuple[Optional[List[str]], List[List[str]]]:
    if not lines:
        return None, []
    header = [c.strip() for c in lines[0].split("|") if c.strip()]
    rows = []
    for line in lines[1:]:
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            rows.append(cells)
    return header if header else None, rows


def _find_metadata_table(tables):
    for i, (header, rows) in enumerate(tables):
        if any("campo" in h.lower() for h in header):
            return i
    return None


def _find_qa_table(tables):
    for i in range(len(tables) - 1, -1, -1):
        header, rows = tables[i]
        if _is_internal_validation_table(header, rows):
            return i
    return None


def _is_internal_validation_table(header, rows) -> bool:
    text = " ".join(header + [cell for row in rows for cell in row])
    text = _normalize_header_cell(text)
    validation_markers = [
        "criterio cumple observacion",
        "usa solo informacion del guion maestro",
        "respeta cantidad y extension del nivel",
        "incluye conexion con ruta",
        "no inventa fuentes ni datos",
        "esta listo para diseno",
    ]
    return any(marker in text for marker in validation_markers)


def _setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    font.color.rgb = COLOR_DARK_GRAY
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)

    for level, size, color in [(1, H1_SIZE, COLOR_NAVY), (2, H2_SIZE, COLOR_BLUE), (3, H3_SIZE, COLOR_NAVY)]:
        heading_style = doc.styles[f"Heading {level}"]
        hfont = heading_style.font
        hfont.name = FONT_NAME
        hfont.size = size
        hfont.color.rgb = color
        hfont.bold = True
        hp = heading_style.paragraph_format
        hp.space_before = Pt(12) if level == 1 else Pt(8)
        hp.space_after = Pt(4)
        hp.line_spacing = 1.15


def _add_cover_page(doc, material_nombre, granule_code, tema, *, program: str, subject: str, level: str):
    _add_styled_para(doc, "", space_before=60, space_after=0)
    _add_styled_para(doc, material_nombre.replace("_", " "),
                      size=COVER_TITLE_SIZE, bold=True, color=COLOR_NAVY,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=10)
    _add_border_below(doc.paragraphs[-1], color_hex=HEX_SEPARATOR, sz="8")
    _add_styled_para(doc, f"{granule_code} \u2014 {tema.replace('-', ' ').title()}",
                      size=COVER_SUBTITLE_SIZE, color=COLOR_BLUE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=14, space_after=6)
    _add_styled_para(doc, f"Programa: {program}",
                      size=COVER_META_SIZE, italic=True, color=COLOR_MED_GRAY,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
    _add_styled_para(doc, f"Asignatura: {subject}",
                      size=COVER_META_SIZE, italic=True, color=COLOR_MED_GRAY,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
    _add_styled_para(doc, f"Nivel: {level}",
                      size=COVER_META_SIZE, italic=True, color=COLOR_MED_GRAY,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=40)
    doc.add_page_break()


def _add_metadata_inline(doc, table_rows):
    pass


def _add_qa_table(doc, table_rows):
    if not table_rows:
        return
    _add_styled_para(doc, "Verificacion interna", size=H3_SIZE,
                      bold=True, color=COLOR_MED_GRAY, space_before=20, space_after=4)
    num_cols = max(len(r) for r in table_rows) if table_rows else 3
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.style = "Table Grid"
    for ri, row in enumerate(table_rows):
        for ci, text in enumerate(row):
            if ci < num_cols:
                cell = table.rows[ri].cells[ci]
                cell.text = text
                if ri == 0:
                    _set_cell_shading(cell, HEX_ACCENT)
                    _set_cell_font(cell, TABLE_HEADER_SIZE, bold=True, color=COLOR_NAVY)
                else:
                    _set_cell_font(cell, BODY_SMALL, color=COLOR_MED_GRAY)
    for ri in range(len(table_rows)):
        for ci in range(num_cols):
            _set_cell_borders(table.rows[ri].cells[ci],
                top={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
                bottom={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
                start={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
                end={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
            )


def _render_fichas(doc, content_tables):
    if not content_tables:
        return
    header, rows = content_tables[0]
    col_map = _header_map(header)

    ficha_idx = 0
    for row in rows:
        ficha_idx += 1
        titulo_a = row[col_map.get("titulo visible lado a", 1)] if len(row) > 1 else ""
        texto_a = row[col_map.get("texto lado a", 2)] if len(row) > 2 else ""
        titulo_b = row[col_map.get("titulo visible lado b", 3)] if len(row) > 3 else ""
        texto_b = row[col_map.get("texto lado b", 4)] if len(row) > 4 else ""
        fuente = row[col_map.get("fuente corta", 5)] if len(row) > 5 else ""

        if ficha_idx > 1:
            _add_separator(doc)

        _add_styled_para(doc, f"Ficha {ficha_idx}: {titulo_a}",
                          size=H2_SIZE, bold=True, color=COLOR_NAVY, space_before=10, space_after=4)

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        cells_data = [
            ("LADO A", texto_a),
            ("LADO B", texto_b),
        ]
        for ri, (label, text) in enumerate(cells_data):
            label_cell = table.rows[ri].cells[0]
            label_cell.text = label
            _set_cell_shading(label_cell, HEX_NAVY)
            _set_cell_font(label_cell, TABLE_HEADER_SIZE, bold=True, color=COLOR_WHITE)
            label_cell.width = Cm(3)

            text_cell = table.rows[ri].cells[1]
            text_cell.text = text
            _set_cell_font(text_cell, TABLE_BODY_SIZE)
            text_cell.width = Cm(12)

        if fuente:
            _add_styled_para(doc, f"Fuente: {fuente}", size=BODY_SMALL,
                              italic=True, color=COLOR_LIGHT_GRAY, space_before=2, space_after=6)


def _render_glosario(doc, content_tables):
    if not content_tables:
        return
    header, rows = content_tables[0]
    col_map = _header_map(header)

    num_cols = 3
    table = doc.add_table(rows=len(rows) + 1, cols=num_cols)
    table.style = "Table Grid"

    hdr_cells = ["No.", "Termino", "Definicion y Aplicacion"]
    for ci, h in enumerate(hdr_cells):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, HEX_NAVY)
        _set_cell_font(cell, TABLE_HEADER_SIZE, bold=True, color=COLOR_WHITE)

    for ri, row in enumerate(rows):
        no = row[col_map.get("no.", 0)] if len(row) > 0 else str(ri + 1)
        termino = row[col_map.get("termino", 1)] if len(row) > 1 else ""
        definicion = row[col_map.get("definicion", 2)] if len(row) > 2 else ""
        aplicacion = row[col_map.get("aplicacion", 4)] if len(row) > 4 else ""
        fuente = row[col_map.get("fuente corta", 5)] if len(row) > 5 else ""

        table.rows[ri + 1].cells[0].text = no
        _set_cell_font(table.rows[ri + 1].cells[0], BODY_SMALL, color=COLOR_MED_GRAY)

        term_cell = table.rows[ri + 1].cells[1]
        term_cell.text = termino
        _set_cell_font(term_cell, BODY_SIZE, bold=True, color=COLOR_NAVY)

        def_cell = table.rows[ri + 1].cells[2]
        def_text = definicion
        if aplicacion:
            def_text += f"\nAplicacion: {aplicacion}"
        if fuente:
            def_text += f"\nFuente: {fuente}"
        def_cell.text = def_text
        _set_cell_font(def_cell, BODY_SMALL)

        if ri % 2 == 1:
            for ci in range(num_cols):
                _set_cell_shading(table.rows[ri + 1].cells[ci], HEX_VERY_LIGHT)

    table.columns[0].width = Cm(1.5)
    table.columns[1].width = Cm(4)
    table.columns[2].width = Cm(10.5)

    for ri in range(len(table.rows)):
        for ci in range(num_cols):
            _set_cell_borders(table.rows[ri].cells[ci],
                top={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
                bottom={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
                start={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
                end={"val": "single", "sz": "2", "color": HEX_SEPARATOR},
            )


def _render_revista(doc, content_tables):
    if not content_tables:
        return
    header, rows = content_tables[0]
    col_map = _header_map(header)

    recuadro_titles = ["evidencia", "criterio tecnico", "riesgo", "decision profesional"]
    ref_titles = ["referencia", "referencias", "uso en revista"]

    for row in rows:
        titulo = row[col_map.get("titulo visible", 1)] if len(row) > 1 else ""
        texto = row[col_map.get("texto final", 3)] if len(row) > 3 else ""
        fuente = row[col_map.get("fuente/seccion del guion maestro", 4)] if len(row) > 4 else ""

        if not titulo and not texto:
            continue

        titulo_lower = titulo.lower().strip()
        is_recuadro = any(r in titulo_lower for r in recuadro_titles)
        is_ref = any(r in titulo_lower for r in ref_titles)

        if is_recuadro:
            _add_callout_box(doc, titulo, texto)
            _set_paragraph_spacing(doc.tables[-1].rows[0].cells[0].paragraphs[0], before=10, after=10)
        elif is_ref:
            _add_styled_para(doc, "Referencias", size=H1_SIZE, bold=True,
                              color=COLOR_NAVY, space_before=18, space_after=8)
            _add_border_below(doc.paragraphs[-1], color_hex=HEX_SEPARATOR, sz="6")
            if texto:
                refs = [r.strip() for r in texto.split(";") if r.strip()]
                if len(refs) == 1 and texto.count("\n") > 0:
                    refs = [r.strip() for r in texto.split("\n") if r.strip()]
                for ref in refs:
                    if ref:
                        _add_bullet_para(doc, ref, size=BODY_SMALL, color=COLOR_DARK_GRAY)
        else:
            _add_styled_para(doc, titulo, size=H1_SIZE if titulo_lower in ["portada", "introduccion"] else H2_SIZE,
                              bold=True, color=COLOR_NAVY, space_before=16, space_after=6)
            if texto:
                _add_text_block(doc, texto, size=BODY_SIZE, color=COLOR_DARK_GRAY,
                                space_before=0, space_after=4)
            if fuente:
                _add_styled_para(doc, f"Fuente: {fuente}", size=BODY_SMALL,
                                  italic=True, color=COLOR_LIGHT_GRAY, space_before=2, space_after=8)


def _render_infografia(doc, content_tables):
    if not content_tables:
        return
    header, rows = content_tables[0]
    col_map = _header_map(header)

    block_num = 0
    for row in rows:
        titulo = row[col_map.get("titulo visible", 1)] if len(row) > 1 else ""
        texto = row[col_map.get("texto visible", 2)] if len(row) > 2 else ""

        if not titulo and not texto:
            continue

        block_num += 1
        block_label = f"Bloque {block_num:02d}"

        block_table = doc.add_table(rows=1, cols=1)
        block_table.style = "Table Grid"
        cell = block_table.rows[0].cells[0]
        cell.text = ""
        _set_cell_shading(cell, HEX_VERY_LIGHT)
        _set_cell_borders(cell,
            top={"val": "single", "sz": "4", "color": HEX_SEPARATOR},
            bottom={"val": "single", "sz": "4", "color": HEX_SEPARATOR},
            start={"val": "single", "sz": "4", "color": HEX_SEPARATOR},
            end={"val": "single", "sz": "4", "color": HEX_SEPARATOR},
        )

        p = cell.paragraphs[0]
        run = p.add_run(f"{block_label} \u2014 {titulo}")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        run.font.color.rgb = COLOR_NAVY
        _set_paragraph_spacing(p, before=6, after=4)

        if texto:
            p2 = cell.add_paragraph()
            run2 = p2.add_run(texto)
            run2.font.name = FONT_NAME
            run2.font.size = BODY_SMALL
            run2.font.color.rgb = COLOR_DARK_GRAY
            _set_paragraph_spacing(p2, before=2, after=6)


def _render_podcast(doc, content_tables):
    if not content_tables:
        return
    header, rows = content_tables[0]
    col_map = _header_map(header)

    first_segment = True
    for row in rows:
        segmento = _row_value(row, col_map, ["segmento"], 0)
        duracion = _row_value(row, col_map, ["duracion estimada", "duracion"], 1)
        texto = _row_value(row, col_map, ["texto de locucion o dialogo", "texto de locución o diálogo", "texto"], 2)
        rol = _row_value(row, col_map, ["rol de voz", "rol"], 3)

        if not segmento and not texto:
            continue

        if not first_segment:
            _add_separator(doc)
        first_segment = False

        _add_styled_para(doc, segmento.upper(), size=H2_SIZE, bold=True,
                          color=COLOR_NAVY, space_before=14, space_after=2)
        if duracion:
            _add_styled_para(doc, f"Duracion: {duracion}", size=BODY_SMALL,
                              italic=True, color=COLOR_LIGHT_GRAY, space_before=0, space_after=6)

        if rol:
            _add_styled_para(doc, f"[{rol.upper()}]", size=BODY_SIZE,
                              bold=True, color=COLOR_BLUE, space_before=6, space_after=2)
        if texto:
            _add_text_block(doc, texto, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=8)


def _render_video(doc, content_tables):
    if not content_tables:
        return
    header, rows = content_tables[0]
    col_map = _header_map(header)

    for row in rows:
        escena_num = _row_value(row, col_map, ["escena"], 0)
        duracion = _row_value(row, col_map, ["duracion", "duracion estimada", "tiempo estimado"], 1)
        funcion = _row_value(row, col_map, ["funcion", "función", "objetivo narrativo"], 2)
        locucion = _row_value(row, col_map, [
            "locucion",
            "locución",
            "guion hablado completo",
            "texto a camara para presentadora",
            "texto a cámara para presentadora",
        ], 3)
        visual = _row_value(row, col_map, ["visual sugerido", "apoyo visual sugerido", "accion visual sugerida", "acción visual sugerida", "visual"], 4)
        recursos = _row_value(row, col_map, ["recursos visuales", "recursos sugeridos", "apoyo audiovisual"])
        texto_pantalla = _row_value(row, col_map, ["texto en pantalla"], 5)
        transicion = _row_value(row, col_map, ["transicion", "transición"])

        if not escena_num and not funcion:
            continue

        escena_label = f"ESCENA {escena_num}" if escena_num else "ESCENA"
        if funcion:
            escena_label += f" \u2014 {funcion}"

        _add_styled_para(doc, escena_label, size=H2_SIZE, bold=True,
                          color=COLOR_NAVY, space_before=14, space_after=2)
        if duracion:
            _add_styled_para(doc, f"Duracion: {duracion}", size=BODY_SMALL,
                              italic=True, color=COLOR_LIGHT_GRAY, space_before=0, space_after=6)

        if locucion:
            _add_styled_para(doc, "LOCUCION:", size=BODY_SMALL,
                              bold=True, color=COLOR_BLUE, space_before=6, space_after=2)
            _add_text_block(doc, locucion, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=4)
        if visual:
            _add_styled_para(doc, "VISUAL:", size=BODY_SMALL,
                              bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
            _add_text_block(doc, visual, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=4)
        if recursos:
            _add_styled_para(doc, "RECURSOS VISUALES:", size=BODY_SMALL,
                              bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
            _add_text_block(doc, recursos, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=4)
        if texto_pantalla:
            _add_styled_para(doc, "TEXTO EN PANTALLA:", size=BODY_SMALL,
                              bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
            _add_text_block(doc, texto_pantalla, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=4)
        if transicion:
            _add_styled_para(doc, "TRANSICION:", size=BODY_SMALL,
                              bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
            _add_text_block(doc, transicion, size=BODY_SIZE,
                            color=COLOR_DARK_GRAY, space_before=0, space_after=4)

        _add_separator(doc)


def _render_video(doc, content_tables):
    if not content_tables:
        return

    for header, rows in content_tables:
        col_map = _header_map(header)
        if "escena" not in col_map:
            continue

        for row in rows:
            escena_num = _row_value(row, col_map, ["escena"], 0)
            duracion = _row_value(row, col_map, ["duracion", "duracion estimada", "tiempo estimado"], 1)
            funcion = _row_value(row, col_map, ["funcion", "función", "objetivo narrativo"], 2)
            locucion = _row_value(row, col_map, [
                "locucion",
                "locución",
                "guion hablado completo",
                "texto a camara para presentadora",
                "texto a cámara para presentadora",
            ], 3)
            visual = _row_value(row, col_map, [
                "visual sugerido",
                "apoyo visual sugerido",
                "accion visual sugerida",
                "acción visual sugerida",
                "visual",
            ], 4)
            recursos = _row_value(row, col_map, ["recursos visuales", "recursos sugeridos", "apoyo audiovisual"])
            texto_pantalla = _row_value(row, col_map, ["texto en pantalla"], 5)
            transicion = _row_value(row, col_map, ["transicion", "transición"])

            if not escena_num and not funcion:
                continue

            escena_label = f"ESCENA {escena_num}" if escena_num else "ESCENA"
            if funcion:
                escena_label += f" - {funcion}"

            _add_styled_para(doc, escena_label, size=H2_SIZE, bold=True,
                              color=COLOR_NAVY, space_before=14, space_after=2)
            if duracion:
                _add_styled_para(doc, f"Duracion: {duracion}", size=BODY_SMALL,
                                  italic=True, color=COLOR_LIGHT_GRAY, space_before=0, space_after=6)
            if locucion:
                _add_styled_para(doc, "LOCUCION:", size=BODY_SMALL,
                                  bold=True, color=COLOR_BLUE, space_before=6, space_after=2)
                _add_text_block(doc, locucion, size=BODY_SIZE,
                                color=COLOR_DARK_GRAY, space_before=0, space_after=4)
            if visual:
                _add_styled_para(doc, "VISUAL:", size=BODY_SMALL,
                                  bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
                _add_text_block(doc, visual, size=BODY_SIZE,
                                color=COLOR_DARK_GRAY, space_before=0, space_after=4)
            if recursos:
                _add_styled_para(doc, "RECURSOS VISUALES:", size=BODY_SMALL,
                                  bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
                _add_text_block(doc, recursos, size=BODY_SIZE,
                                color=COLOR_DARK_GRAY, space_before=0, space_after=4)
            if texto_pantalla:
                _add_styled_para(doc, "TEXTO EN PANTALLA:", size=BODY_SMALL,
                                  bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
                _add_text_block(doc, texto_pantalla, size=BODY_SIZE,
                                color=COLOR_DARK_GRAY, space_before=0, space_after=4)
            if transicion:
                _add_styled_para(doc, "TRANSICION:", size=BODY_SMALL,
                                  bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
                _add_text_block(doc, transicion, size=BODY_SIZE,
                                color=COLOR_DARK_GRAY, space_before=0, space_after=4)

            _add_separator(doc)


def _add_footer_with_page_number(doc, granule_code, material_nombre):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    while len(footer.paragraphs) > 1:
        footer.paragraphs[-1]._element.getparent().remove(footer.paragraphs[-1]._element)

    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"{granule_code}  \u2014  {material_nombre.replace('_', ' ')}")
    run.font.size = FOOTER_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = COLOR_LIGHT_GRAY

    run = p.add_run("  \u00b7  ")
    run.font.size = FOOTER_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = COLOR_LIGHT_GRAY

    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = FOOTER_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = COLOR_LIGHT_GRAY

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _extract_video_scene_rows(content_tables):
    scenes = []
    for header, rows in content_tables:
        col_map = _header_map(header)
        if "escena" not in col_map:
            continue
        for row in rows:
            escena_num = _row_value(row, col_map, ["escena"], 0)
            funcion = _row_value(row, col_map, ["funcion", "función", "objetivo narrativo"], 2)
            if not escena_num and not funcion:
                continue
            scenes.append({
                "escena": escena_num,
                "duracion": _row_value(row, col_map, ["duracion", "duracion estimada", "tiempo estimado"], 1),
                "funcion": funcion,
                "locucion": _row_value(row, col_map, [
                    "locucion",
                    "locución",
                    "guion hablado completo",
                    "texto a camara para presentadora",
                    "texto a cámara para presentadora",
                ], 3),
                "visual": _row_value(row, col_map, [
                    "visual sugerido",
                    "apoyo visual sugerido",
                    "accion visual sugerida",
                    "acción visual sugerida",
                    "visual",
                ], 4),
                "recursos": _row_value(row, col_map, ["recursos visuales", "recursos sugeridos", "apoyo audiovisual"]),
                "texto_pantalla": _row_value(row, col_map, ["texto en pantalla"], 5),
                "transicion": _row_value(row, col_map, ["transicion", "transición"]),
            })
    return scenes


def _add_video_scene_table(doc, scenes):
    if not scenes:
        return
    _add_styled_para(doc, "Tabla de escenas", size=H1_SIZE, bold=True,
                      color=COLOR_NAVY, space_before=8, space_after=6)
    headers = ["Escena", "Tiempo", "Objetivo", "Guion hablado", "Visual", "Recursos", "Texto en pantalla", "Transicion"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        _set_cell_shading(cell, HEX_ACCENT)
        _set_cell_font(cell, TABLE_HEADER_SIZE, bold=True, color=COLOR_NAVY)
    for scene in scenes:
        cells = table.add_row().cells
        values = [
            scene["escena"],
            scene["duracion"],
            scene["funcion"],
            scene["locucion"],
            scene["visual"],
            scene["recursos"],
            scene["texto_pantalla"],
            scene["transicion"],
        ]
        for idx, value in enumerate(values):
            cells[idx].text = _clean_visible_text(value)
            _set_cell_font(cells[idx], BODY_SMALL, color=COLOR_DARK_GRAY)
    _add_separator(doc)


def _render_video(doc, content_tables):
    scenes = _extract_video_scene_rows(content_tables)
    if not scenes:
        return

    _add_video_scene_table(doc, scenes)
    _add_styled_para(doc, "Desarrollo por escena", size=H1_SIZE, bold=True,
                      color=COLOR_NAVY, space_before=12, space_after=6)
    for scene in scenes:
        escena_label = f"ESCENA {scene['escena']}" if scene["escena"] else "ESCENA"
        if scene["funcion"]:
            escena_label += f" - {scene['funcion']}"
        _add_styled_para(doc, escena_label, size=H2_SIZE, bold=True,
                          color=COLOR_NAVY, space_before=14, space_after=2)
        if scene["duracion"]:
            _add_styled_para(doc, f"Duracion: {scene['duracion']}", size=BODY_SMALL,
                              italic=True, color=COLOR_LIGHT_GRAY, space_before=0, space_after=6)
        for label, key in [
            ("LOCUCION:", "locucion"),
            ("VISUAL:", "visual"),
            ("RECURSOS VISUALES:", "recursos"),
            ("TEXTO EN PANTALLA:", "texto_pantalla"),
            ("TRANSICION:", "transicion"),
        ]:
            if scene[key]:
                _add_styled_para(doc, label, size=BODY_SMALL,
                                  bold=True, color=COLOR_BLUE, space_before=4, space_after=2)
                _add_text_block(doc, scene[key], size=BODY_SIZE,
                                color=COLOR_DARK_GRAY, space_before=0, space_after=4)
        _add_separator(doc)


RENDERERS = {
    "02": _render_fichas,
    "03": _render_glosario,
    "04": _render_revista,
    "05": _render_infografia,
    "06": _render_podcast,
    "07": _render_video,
}

# Claves internas de layout ("02"…"07"): mismas que MATERIAL_VALIDATION_RULES / RENDERERS.
# Por categoría académica, el nn del material (01…07 en prompts) no coincide con esa clave:
# hay que mapear qué plantilla DOCX usar (fichas, glosario, revista, infografía, podcast, video).
_LAYOUT_RENDERER_BY_CATEGORY: dict[str, dict[str, str]] = {
    "especializacion": {"02": "02", "03": "03", "04": "04", "05": "05", "06": "06", "07": "07"},
    "pregrado": {
        "01": "06",
        "02": "05",
        "03": "07",
        "04": "03",
        "05": "07",
        "06": "04",
        "07": "02",
    },
    "curso_rapido": {
        "01": "05",
        "02": "06",
        "03": "07",
        "04": "03",
        "05": "07",
        "06": "04",
        "07": "02",
    },
    "diplomado": {
        "01": "07",
        "02": "03",
        "03": "04",
        "04": "07",
        "05": "05",
        "06": "06",
        "07": "02",
    },
    "curso_externos_profesional": {
        "01": "07",
        "02": "06",
        "03": "05",
        "04": "07",
        "05": "04",
        "06": "03",
        "07": "02",
    },
}


MATERIAL_NUMBER_BY_NAME = {
    rule["name"]: nn
    for nn, rule in MATERIAL_VALIDATION_RULES.items()
}


def _get_material_number(material_nombre: str) -> Optional[str]:
    return MATERIAL_NUMBER_BY_NAME.get(material_nombre)


def resolve_layout_renderer_key(
    category_key: str | None,
    material_nn: str | None,
    material_nombre: str,
) -> Optional[str]:
    """Resuelve la clave de RENDERERS ("02"…"07") según categoría y número de material."""
    ck = (category_key or "").strip().lower()
    nn = (material_nn or "").strip()
    table = _LAYOUT_RENDERER_BY_CATEGORY.get(ck)
    if table and nn in table:
        return table[nn]
    return _get_material_number(material_nombre)


def _validate_rendered_docx(
    output_path: Path,
    nn: Optional[str],
    category_key: Optional[str] = None,
    material_nn: Optional[str] = None,
    material_nombre: Optional[str] = None,
) -> None:
    doc = Document(output_path)
    body_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            body_parts.append(" | ".join(cell.text for cell in row.cells))
    body_text = "\n".join(body_parts)
    body_lower = body_text.lower()

    forbidden = [
        "<ul", "</ul", "<li", "</li", "<ol", "</ol",
        "verificacion interna", "verificación interna",
        "verificacion de calidad", "verificación de calidad",
        "checklist",
        "usa solo informacion del guion maestro", "usa solo información del guion maestro",
        "respeta cantidad y extension del nivel", "respeta cantidad y extensión del nivel",
        "no inventa fuentes ni datos", "esta listo para diseno", "está listo para diseño",
    ]
    found = [item for item in forbidden if item in body_lower]
    if found:
        raise ValueError(f"DOCX contiene elementos internos o HTML visible: {', '.join(found)}")

    minimals = _get_validation_minimals(category_key, nn, material_nn)
    mat_type = minimals.get("type", "")
    material_ref = material_nombre or material_nn or nn

    if mat_type == "fichas":
        min_fichas = minimals.get("min_fichas", 5)
        ficha_count = sum(1 for p in body_parts if p.strip().lower().startswith("ficha "))
        if ficha_count < min_fichas:
            raise ValueError(f"DOCX {material_ref} incompleto: {ficha_count}/{min_fichas} fichas renderizadas. Categoria: {category_key}, material: {material_nn}, layout: {nn}")
    elif mat_type == "glosario":
        min_terminos = minimals.get("min_terminos", 12)
        if not doc.tables or len(doc.tables[0].rows) < min_terminos + 1:
            rows = len(doc.tables[0].rows) if doc.tables else 0
            raise ValueError(f"DOCX {material_ref} incompleto: {max(0, rows - 1)}/{min_terminos} terminos renderizados. Categoria: {category_key}, material: {material_nn}, layout: {nn}")
    elif mat_type == "revista":
        min_bloques = minimals.get("min_bloques", 14)
        if len(body_parts) < min_bloques + 4:
            raise ValueError(f"DOCX {material_ref} incompleto: solo {len(body_parts)} bloques/parrafos visibles (min {min_bloques + 4}). Categoria: {category_key}, material: {material_nn}, layout: {nn}")
    elif mat_type == "infografia":
        min_bloques = minimals.get("min_bloques", 7)
        block_count = sum(
            1 for table in doc.tables
            if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip().lower().startswith("bloque ")
        )
        if block_count < min_bloques:
            raise ValueError(f"DOCX {material_ref} incompleto: {block_count}/{min_bloques} bloques renderizados. Categoria: {category_key}, material: {material_nn}, layout: {nn}")
    elif mat_type == "podcast":
        min_segmentos = minimals.get("min_segmentos", 9)
        segment_count = sum(1 for p in body_parts if p.strip().lower().startswith("duracion:"))
        if segment_count < min_segmentos:
            raise ValueError(f"DOCX {material_ref} incompleto: {segment_count}/{min_segmentos} segmentos renderizados. Categoria: {category_key}, material: {material_nn}, layout: {nn}")
    elif mat_type == "video":
        min_escenas = minimals.get("min_escenas", 7)
        scene_count = sum(1 for p in body_parts if p.strip().lower().startswith("escena"))
        if scene_count < min_escenas:
            raise ValueError(f"DOCX {material_ref} incompleto: {scene_count}/{min_escenas} escenas renderizadas. Categoria: {category_key}, material: {material_nn}, layout: {nn}")
        if len(body_text) < 2500:
            raise ValueError(f"DOCX {material_ref} incompleto: solo {len(body_text)} caracteres visibles. Categoria: {category_key}, material: {material_nn}, layout: {nn}")


def save_docx_with_structure(
    content: str,
    output_path: Path,
    material_nombre: str,
    granule_code: str,
    tema: str,
    *,
    category_key: str | None = None,
    material_nn: str | None = None,
    program: str = "",
    subject: str = "",
    level: str = "",
) -> None:
    cleaned = clean_ai_response(content)
    cover_metadata = f"Programa: {program}\nAsignatura: {subject}\nNivel: {level}"
    print("[Materials][CoverMetadata]")
    print(f"program={program}")
    print(f"subject={subject}")
    print(f"level={level}")
    if not (program or "").strip() or not (subject or "").strip() or not (level or "").strip():
        raise ValueError(
            "Metadata de portada incompleta: "
            f"program={program!r}, subject={subject!r}, level={level!r}. "
            "Se aborta antes de renderizar."
        )
    resource = f"{granule_code} {material_nombre}"
    _assert_no_contaminated_metadata(content, "raw_content", resource)
    _assert_no_contaminated_metadata(cleaned, "cleaned_content", resource)
    _assert_no_contaminated_metadata(cover_metadata, "cover metadata", resource)
    doc = Document()
    _setup_document(doc)
    _add_cover_page(doc, material_nombre, granule_code, tema, program=program, subject=subject, level=level)
    chars_after_cover = count_visible_docx_chars(doc)

    tables = _parse_markdown_tables(cleaned)
    meta_idx = _find_metadata_table(tables)
    qa_idx = _find_qa_table(tables)

    content_tables = []
    for i, (header, rows) in enumerate(tables):
        if i == meta_idx or i == qa_idx:
            continue
        if _is_internal_validation_table(header, rows):
            continue
        content_tables.append((header, rows))

    if meta_idx is not None:
        _add_metadata_inline(doc, tables[meta_idx][1])

    layout_nn = resolve_layout_renderer_key(category_key, material_nn, material_nombre)
    if not layout_nn:
        raise ValueError(
            f"No se pudo resolver plantilla DOCX para material {material_nombre!r} "
            f"(categoría={category_key!r}, nn={material_nn!r}). "
            "Compruebe resolve_layout_renderer_key / categoría en configuración."
        )
    renderer = RENDERERS.get(layout_nn)
    if not renderer:
        raise ValueError(f"Plantilla de renderizado interna ausente para layout {layout_nn!r}")

    renderer(doc, content_tables)

    chars_after_render = count_visible_docx_chars(doc)
    if chars_after_render <= chars_after_cover + 200:
        print(
            "    ADVERTENCIA: renderer sin contenido suficiente; "
            "se inserta fallback con respuesta limpia completa."
        )
        _add_raw_content_fallback(doc, cleaned)

    _add_footer_with_page_number(doc, granule_code, material_nombre)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    _validate_rendered_docx(output_path, layout_nn, category_key, material_nn, material_nombre)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells if (cell.text or "").strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_especializacion_prompt() -> str:
    if not ESPECIALIZACION_PROMPT_PATH.exists():
        raise FileNotFoundError(
            f"No se encontro el prompt de materiales: {ESPECIALIZACION_PROMPT_PATH}. "
            "Verifica que el archivo existe en prompts/."
        )
    return ESPECIALIZACION_PROMPT_PATH.read_text(encoding="utf-8")


def extract_system_prompt(prompt_text: str) -> str:
    match = re.search(r"# PROMPT SISTEMA\s*\n\s*```text\s*\n(.*?)```", prompt_text, re.DOTALL)
    if match:
        return match.group(1).strip()
    raise ValueError(
        "No se encontro la seccion '# PROMPT SISTEMA' con bloque ```text``` en el archivo de prompt."
    )


def extract_material_prompt(prompt_text: str, seccion_nombre: str) -> str:
    pattern = rf"## {re.escape(seccion_nombre)}\s*\n.*?### Prompt\s*\n\s*```text\s*\n(.*?)```"
    match = re.search(pattern, prompt_text, re.DOTALL)
    if match:
        return match.group(1).strip()
    raise ValueError(
        f"No se encontro la seccion de prompt para '{seccion_nombre}'."
    )


def build_user_prompt(
    material: MaterialConfig,
    guion_maestro_text: str,
    granule_code: str,
    tema: str,
    tema_corto: str,
    version: str,
    program: str,
    subject: str,
    level: str,
) -> str:
    return f"""Programa oficial: {program}
Asignatura oficial: {subject}
Nivel oficial: {level}
Granulo oficial: {granule_code} - {tema}

Esta prohibido mencionar cualquier programa, asignatura o especializacion distinta.

Quiero generar un material derivado para {level.upper()}.

Pego a continuacion el GUION MAESTRO aprobado del tema:

{guion_maestro_text}

Datos del material:
- Codigo GX: {granule_code}
- Nombre exacto del tema: {tema}
- Nombre corto para archivo: {tema_corto}
- Version: {version}
- Material a generar: {material.nombre.replace("_", " ")}
- Formato esperado: DOCX (contenido en tabla)
- Cierre integrado ira en: NO APLICA

Instrucciones especificas para este material:

{material.prompt_particular}

REGLAS CRITICAS DE SALIDA:
1. NO incluyas la frase "Datos recibidos" ni ninguna frase de confirmacion de instrucciones.
2. NO confirmes que entendiste las instrucciones. Entrega directamente el contenido final.
3. NO uses markdown fences (```text, ```, ```markdown).
4. NO inventes fuentes, terminos, casos, tecnologias ni datos que no esten en el GUION MAESTRO.
5. Si una informacion no esta en el guion maestro, marca como "Informacion faltante" en una tabla.
6. Entrega unicamente el contenido del material solicitado en formato tabla markdown.
7. No generes los demas materiales. Solo este.

REGLAS EDITORIALES MINIMAS:
1. Aunque la estructura tecnica se entregue en tabla markdown para validacion, redacta cada celda con lenguaje editorial profesional.
2. Evita fragmentos frios o telegraficos: usa parrafos naturales, legibles y listos para maquetacion academica.
3. Revista, Podcast y Video deben sentirse como documentos profesionales reales desde su redaccion base.
4. Las referencias, fuentes y conexiones de ruta deben mantenerse claras, pero sin saturar el texto principal.
5. Si el material incluye un bloque llamado "Conceptos clave", no lo resumas como glosario. Desarrollalo como seccion editorial con apertura general y conceptos explicados en profundidad, respetando la extension indicada en el prompt particular.

Genera unicamente el material solicitado.
""".strip()


def clean_ai_response(content: str) -> str:
    cleaned = content.strip()
    for pattern in ARTIFACT_PATTERNS:
        cleaned = pattern.sub("", cleaned)
    lines = [line for line in cleaned.splitlines() if line.strip()]
    cleaned = "\n".join(lines).strip()
    return cleaned


def _find_plan_curso(generated_dir: Path, output_base: Path) -> Path | None:
    candidates = [
        generated_dir / "plan_curso.json",
        generated_dir.parent / "plan_curso.json",
        output_base / "plan_curso.json",
        output_base.parent / "plan_curso.json",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _load_required_plan_metadata(generated_dir: Path, output_base: Path) -> dict:
    plan_path = _find_plan_curso(generated_dir, output_base)
    if not plan_path:
        searched = [
            generated_dir / "plan_curso.json",
            generated_dir.parent / "plan_curso.json",
            output_base / "plan_curso.json",
            output_base.parent / "plan_curso.json",
        ]
        raise FileNotFoundError(
            "plan_curso.json es obligatorio para generar materiales. Buscado en: "
            + "; ".join(str(path) for path in searched)
        )
    data = json.loads(plan_path.read_text(encoding="utf-8"))
    print(f"plan keys={sorted(data.keys())}")
    return {
        "path": str(plan_path),
        "program": data.get("programa", ""),
        "subject": data.get("asignatura", ""),
        "level": data.get("nivel", data.get("categoria", "")),
        "topics": data.get("temas", []),
    }


def _normalize_metadata_check(value: str) -> str:
    value = (value or "").lower()
    replacements = str.maketrans("áéíóúüñ", "aeiouun")
    value = value.translate(replacements)
    value = re.sub(r"\s+", " ", value)
    return value


def _has_contaminated_metadata(value: str) -> bool:
    normalized = _normalize_metadata_check(value)
    return "videojuegos" in normalized or all(token in normalized for token in ("diseno", "desarrollo", "videojuegos"))


def _contamination_excerpt(value: str) -> str:
    if not value:
        return ""
    normalized = _normalize_metadata_check(value)
    index = normalized.find("videojuegos")
    if index < 0:
        index = normalized.find("diseno")
    if index < 0:
        return ""
    start = max(0, index - 100)
    end = min(len(value), index + 180)
    return re.sub(r"\s+", " ", value[start:end]).strip()


def _assert_no_contaminated_metadata(value: str, source: str, resource: str = "") -> None:
    if _has_contaminated_metadata(value):
        print("[Materials][Contamination]")
        print(f"resource={resource or 'unknown'}")
        print("phrase=metadata ajena de videojuegos")
        print(f"source={source}")
        print(f"excerpt={_contamination_excerpt(value)}")
        raise ValueError(f"Metadata contaminada detectada en {source}; se aborta antes de guardar DOCX.")


def generate_material_content(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int,
    temperature: float,
) -> str:
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return (response.choices[0].message.content or "").strip()


def build_material_filename(nn: str, granule_code: str, material_nombre: str, tema_corto: str, version: str, extension: str) -> str:
    tema_clean = normalize_for_filename(tema_corto)
    material_clean = normalize_for_filename(material_nombre)
    return f"{nn}_{granule_code}_{material_clean}_{tema_clean}_{version}{extension}"


def discover_granules(generated_dir: Path) -> List[dict]:
    if not generated_dir.exists() or not generated_dir.is_dir():
        raise FileNotFoundError(f"El directorio de granulos no existe: {generated_dir}")

    granule_pattern = re.compile(r"^(G[1-5])_(.+)\.docx$", re.IGNORECASE)
    granules = []

    for entry in generated_dir.iterdir():
        if not entry.is_file():
            continue
        if entry.name.startswith("~$"):
            continue
        match = granule_pattern.match(entry.name)
        if match:
            granule_code = match.group(1).upper()
            tema_slug = match.group(2)
            granules.append({
                "code": granule_code,
                "tema": tema_slug,
                "tema_corto": normalize_for_filename(tema_slug),
                "path": entry,
            })

    if not granules:
        raise ValueError(
            f"No se encontraron archivos de granulos (G1_*.docx a G5_*.docx) en {generated_dir}."
        )

    granules.sort(key=lambda g: int(g["code"][1:]))
    return granules


def validate_material_prompts(prompt_text: str) -> List[str]:
    missing = []
    for mat in MATERIALES_A_GENERAR:
        try:
            extract_material_prompt(prompt_text, mat["seccion_prompt"])
        except ValueError:
            missing.append(mat["seccion_prompt"])
    return missing


def validate_material_content(
    nn: str,
    content: str,
    category_key: str | None = None,
    material_nn: str | None = None,
) -> Tuple[str, List[str]]:
    minimals = _get_validation_minimals(category_key, nn, material_nn)
    if not minimals:
        return "ok", []
    warnings = []
    table_rows = 0
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|") and "---" not in stripped:
            table_rows += 1

    mat_type = minimals.get("type", "")
    if mat_type == "fichas":
        expected_rows = minimals.get("min_fichas", 5) + 1
        if table_rows < expected_rows:
            warnings.append(
                f"FICHAS: se esperaban al menos {minimals['min_fichas']} fichas (filas de tabla). "
                f"Se detectaron {max(0, table_rows - 1)} filas de contenido."
            )
    elif mat_type == "glosario":
        expected_rows = minimals.get("min_terminos", 12) + 1
        if table_rows < expected_rows:
            warnings.append(
                f"GLOSARIO: se esperaban al menos {minimals['min_terminos']} terminos. "
                f"Se detectaron {max(0, table_rows - 1)} filas de contenido."
            )
    elif mat_type == "revista":
        expected_rows = minimals.get("min_bloques", 14) + 1
        if table_rows < expected_rows:
            warnings.append(
                f"REVISTA: se esperaban al menos {minimals['min_bloques']} bloques. "
                f"Se detectaron {max(0, table_rows - 1)} filas de contenido."
            )
    elif mat_type == "infografia":
        expected_rows = minimals.get("min_bloques", 7) + 1
        if table_rows < expected_rows:
            warnings.append(
                f"INFOGRAFIA: se esperaban al menos {minimals['min_bloques']} bloques. "
                f"Se detectaron {max(0, table_rows - 1)} filas de contenido."
            )
    elif mat_type == "podcast":
        expected_rows = minimals.get("min_segmentos", 9) + 1
        if table_rows < expected_rows:
            warnings.append(
                f"PODCAST: se esperaban al menos {minimals['min_segmentos']} segmentos. "
                f"Se detectaron {max(0, table_rows - 1)} filas de contenido."
            )
    elif mat_type == "video":
        expected_rows = minimals.get("min_escenas", 7) + 1
        if table_rows < expected_rows:
            warnings.append(
                f"VIDEO: se esperaban al menos {minimals['min_escenas']} escenas. "
                f"Se detectaron {max(0, table_rows - 1)} filas de contenido."
            )

    status = "warning" if warnings else "ok"
    return status, warnings


def generate_all_materiales(
    job_id: str,
    generated_dir: Path,
    output_base: Path,
    model: str,
    max_tokens: int,
    temperature: float,
) -> dict:
    if not model:
        model = get_openai_model("materials")

    print(f"\n{'=' * 60}")
    print(f"=== FASE 3: GENERACION DE MATERIALES DE ESPECIALIZACION ===")
    print(f"{'=' * 60}")
    print(f"Job ID: {job_id}")
    print(f"Directorio de granulos: {generated_dir}")
    print(f"Directorio de salida: {output_base}")
    print(f"")
    print(f"NOTA: En Especializacion se generan 6 materiales por granulo.")
    print(f"Se excluye 01_VIDEO_CASO_O_PROBLEMA (reservado para presentadoras).")
    print(f"Materiales a generar:")
    print(f"  02 - FICHAS DE ESTUDIO DE EVIDENCIA")
    print(f"  03 - GLOSARIO ESPECIALIZADO")
    print(f"  04 - REVISTA DOSSIER")
    print(f"  05 - INFOGRAFIA MODELO O RUTA")
    print(f"  06 - PODCAST DEBATE EXPERTO")
    print(f"  07 - VIDEO SOLUCION O PROCEDIMIENTO")

    prompt_text = load_especializacion_prompt()
    print(f"\nPrompt cargado: {ESPECIALIZACION_PROMPT_PATH}")

    system_prompt = extract_system_prompt(prompt_text)
    print(f"System prompt extraido correctamente.")

    missing_prompts = validate_material_prompts(prompt_text)
    if missing_prompts:
        raise ValueError(
            f"Faltan {len(missing_prompts)} bloques de prompt. "
            f"Secciones no encontradas: {', '.join(missing_prompts)}."
        )
    print(f"Validacion de prompts: {EXPECTED_MATERIAL_COUNT}/{EXPECTED_MATERIAL_COUNT} bloques encontrados.")

    materiales_config: List[MaterialConfig] = []
    for mat in MATERIALES_A_GENERAR:
        try:
            prompt_particular = extract_material_prompt(prompt_text, mat["seccion_prompt"])
            materiales_config.append(MaterialConfig(
                nn=mat["nn"],
                nombre=mat["nombre"],
                prompt_particular=prompt_particular,
            ))
            print(f"  Material configurado: {mat['nn']} - {mat['nombre']}")
        except ValueError as e:
            print(f"  ADVERTENCIA: {e}")

    if len(materiales_config) != EXPECTED_MATERIAL_COUNT:
        raise ValueError(
            f"Se esperaban {EXPECTED_MATERIAL_COUNT} materiales configurados, "
            f"pero se encontraron {len(materiales_config)}."
        )

    granules = discover_granules(generated_dir)
    print(f"\nGranulos encontrados: {len(granules)}")
    for g in granules:
        print(f"  - {g['code']}: {g['tema']}")

    if len(granules) != EXPECTED_GRANULE_COUNT:
        print(
            f"ADVERTENCIA: Se encontraron {len(granules)} granulos, "
            f"se esperaban {EXPECTED_GRANULE_COUNT}."
        )

    print(f"Materiales por granulo: {len(materiales_config)}")
    print(f"Total de materiales a generar: {len(granules) * len(materiales_config)}")

    plan_metadata = _load_required_plan_metadata(generated_dir, output_base)
    print("[Materials][MetadataSource]")
    print(f"program from=plan_curso.json value={plan_metadata.get('program', '')}")
    print(f"subject from=plan_curso.json value={plan_metadata.get('subject', '')}")
    print(f"level from=plan_curso.json value={plan_metadata.get('level', '')}")

    client = get_openai_client()
    errors = []
    manifest_entries = []
    summary = {
        "job_id": job_id,
        "fecha": datetime.now(timezone.utc).isoformat(),
        "total_granules": len(granules),
        "total_materiales_esperados": len(granules) * len(materiales_config),
        "total_materiales_generados": 0,
        "total_errores": 0,
        "total_advertencias": 0,
        "granules": {},
    }

    for granule in granules:
        granule_code = granule["code"]
        tema = granule["tema"]
        tema_corto = granule["tema_corto"]
        granule_path = granule["path"]

        print(f"\n{'-' * 50}")
        print(f"--- Procesando {granule_code}: {tema} ---")
        print(f"{'-' * 50}")

        try:
            guion_text = extract_docx_text(granule_path)
            print(f"  Contenido del granulo leido: {len(guion_text)} caracteres")
        except Exception as e:
            error_msg = f"Error leyendo {granule_path}: {e}"
            print(f"  ERROR: {error_msg}")
            errors.append({"granule": granule_code, "error": error_msg})
            summary["granules"][granule_code] = {"status": "error_lectura", "materiales": {}}
            summary["total_errores"] += 1
            continue

        folder_name = build_granule_folder_name(granule_code, tema)
        granule_output_dir = output_base / folder_name
        granule_output_dir.mkdir(parents=True, exist_ok=True)
        print(f"  Carpeta de salida: {granule_output_dir}")
        print(f"granule from=filename/docx value={granule_code} {tema}")

        granule_summary = {"status": "ok", "materiales": {}}
        granule_errors = []
        granule_warnings_count = 0

        for material in materiales_config:
            material_filename = build_material_filename(
                material.nn, granule_code, material.nombre, tema_corto, VERSION_DEFECTO, ".docx"
            )
            material_output_path = granule_output_dir / material_filename

            print(f"  [{material.nn}/6] {material.nombre}")
            print(f"    Archivo: {material_filename}")

            try:
                user_prompt = build_user_prompt(
                    material=material,
                    guion_maestro_text=guion_text,
                    granule_code=granule_code,
                    tema=tema,
                    tema_corto=tema_corto,
                    version=VERSION_DEFECTO,
                    program=plan_metadata.get("program", ""),
                    subject=plan_metadata.get("subject", ""),
                    level=plan_metadata.get("level", ""),
                )

                client = get_openai_client()
                content = generate_material_content(
                    client=client,
                    model=model,
                    system_prompt=system_prompt,
                    user_prompt=user_prompt,
                    max_tokens=max_tokens,
                    temperature=temperature,
                )

                _assert_no_contaminated_metadata(content, "raw_content", f"{granule_code} {material.nombre}")

                if not content or len(content) < MIN_RESPONSE_CHARS:
                    raise ValueError(
                        f"Respuesta insuficiente ({len(content)} chars). Minimo: {MIN_RESPONSE_CHARS}."
                    )

                layout_nn = resolve_layout_renderer_key("especializacion", material.nn, material.nombre)
                if not layout_nn:
                    raise ValueError(f"No se resolvió layout para especialización material {material.nn} {material.nombre}")
                val_status, val_warnings = validate_material_content(layout_nn, content, "especializacion", material.nn)
                if val_warnings:
                    for w in val_warnings:
                        print(f"    ADVERTENCIA: {w}")
                    granule_warnings_count += len(val_warnings)

                save_docx_with_structure(
                    content=content,
                    output_path=material_output_path,
                    material_nombre=material.nombre,
                    granule_code=granule_code,
                    tema=tema,
                    category_key="especializacion",
                    material_nn=material.nn,
                    program=plan_metadata.get("program", ""),
                    subject=plan_metadata.get("subject", ""),
                    level=plan_metadata.get("level", ""),
                )

                if not material_output_path.exists() or material_output_path.stat().st_size == 0:
                    raise ValueError("El archivo se guardo vacio o no se creo.")

                file_size = material_output_path.stat().st_size
                print(f"    Guardado: {material_filename} ({file_size} bytes)")

                granule_summary["materiales"][material.nn] = {
                    "nombre": material.nombre,
                    "archivo": material_filename,
                    "status": "ok",
                    "validation_status": val_status,
                    "warnings": val_warnings,
                    "size_bytes": file_size,
                }
                summary["total_materiales_generados"] += 1

                manifest_entries.append({
                    "granule_code": granule_code,
                    "granule_topic": tema,
                    "material_number": material.nn,
                    "material_name": material.nombre,
                    "filename": material_filename,
                    "path": str(material_output_path),
                    "validation_status": val_status,
                    "warnings": val_warnings,
                })

            except Exception as e:
                error_msg = f"Error material {material.nn} {granule_code} ({material.nombre}): {e}"
                print(f"    ERROR: {error_msg}")
                granule_errors.append({
                    "granule": granule_code,
                    "material": material.nn,
                    "nombre": material.nombre,
                    "error": str(e),
                    "traceback": traceback.format_exc(),
                })
                granule_summary["materiales"][material.nn] = {
                    "nombre": material.nombre,
                    "status": "error",
                    "error": str(e),
                }
                summary["total_errores"] += 1

        if granule_errors:
            granule_summary["status"] = "parcial"
            errors.extend(granule_errors)
        if granule_warnings_count > 0:
            summary["total_advertencias"] = summary.get("total_advertencias", 0) + granule_warnings_count

        summary["granules"][granule_code] = granule_summary
        print(f"  {granule_code} completado: estado={granule_summary['status']}, "
              f"advertencias={granule_warnings_count}")

    manifest = {
        "job_id": job_id,
        "fecha": datetime.now(timezone.utc).isoformat(),
        "materiales": manifest_entries,
    }

    return {
        "summary": summary,
        "manifest": manifest,
        "errors": errors,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Genera 6 materiales de especializacion derivados de cada granulo (G1-G5)."
    )
    parser.add_argument("--job-id", required=True, help="ID del job")
    parser.add_argument("--generated-dir", required=True, help="Directorio con granulos G1-G5")
    parser.add_argument("--output-dir", required=True, help="Directorio base para materiales")
    parser.add_argument("--model", default=None, help="Modelo (fallback: OPENAI_MODEL_MATERIALS > OPENAI_MODEL > gemini-2.5-flash)")
    parser.add_argument("--max-tokens", type=int, default=8000, help="Maximo de tokens por material")
    parser.add_argument("--temperature", type=float, default=0.5, help="Creatividad de generacion")
    return parser.parse_args()


def main() -> None:
    if load_dotenv:
        load_dotenv()

    args = parse_args()
    generated_dir = Path(args.generated_dir)
    output_base = Path(args.output_dir)
    output_base.mkdir(parents=True, exist_ok=True)

    result = generate_all_materiales(
        job_id=args.job_id,
        generated_dir=generated_dir,
        output_base=output_base,
        model=args.model,
        max_tokens=args.max_tokens,
        temperature=args.temperature,
    )

    summary_path = output_base.parent / "summary.json"
    summary_path.write_text(
        json.dumps(result["summary"], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"\nSummary guardado: {summary_path}")

    manifest_path = output_base.parent / "manifest.json"
    manifest_path.write_text(
        json.dumps(result["manifest"], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"Manifest guardado: {manifest_path}")

    errors_path = output_base.parent / "errors.json"
    errors_path.write_text(
        json.dumps(result["errors"], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"Errors guardado: {errors_path}")

    print(f"\n{'=' * 60}")
    print(f"=== RESUMEN FINAL ===")
    print(f"{'=' * 60}")
    print(f"Job ID: {result['summary']['job_id']}")
    print(f"Granulos procesados: {result['summary']['total_granules']}")
    print(f"Materiales generados: {result['summary']['total_materiales_generados']}")
    print(f"Materiales esperados: {result['summary']['total_materiales_esperados']}")
    print(f"Errores: {result['summary']['total_errores']}")
    print(f"Advertencias: {result['summary'].get('total_advertencias', 0)}")

    if result["errors"]:
        print(f"\n--- Errores detectados ---")
        for err in result["errors"]:
            g = err.get("granule", "?")
            m = err.get("material", "?")
            n = err.get("nombre", "?")
            e = err.get("error", "?")
            print(f"  [{g}] Material {m} ({n}): {e}")
        print(f"--- Fin de errores ---")
    else:
        print(f"\nTodos los materiales se generaron sin errores.")

    print(f"\nGeneracion completa.")


if __name__ == "__main__":
    main()
