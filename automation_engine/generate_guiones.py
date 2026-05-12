import argparse
import json
import os
import re
import sys
import textwrap
import traceback
import unicodedata
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

try:
    from dotenv import load_dotenv
except ImportError:  # pragma: no cover
    load_dotenv = None

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover
    OpenAI = None

from automation_engine.config.categories import CATEGORIES


ENGINE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = ENGINE_DIR.parent
PROMPT_PATHS = {key: config.guion_prompt_path for key, config in CATEGORIES.items()}

DOMINANT_KEYWORD_LIMITS = {
    "contexto latinoamericano": 3,
    "innovacion": 5,
    "innovación": 5,
    "innovaci\ufffdn": 5,
    "inmersion": 4,
    "inmersión": 4,
    "inmersi\ufffdn": 4,
    "narrativa interactiva": 4,
    "narrativas interactivas": 4,
    "accesibilidad": 5,
    "sostenibilidad": 4,
}

DOMINANT_KEYWORD_REPLACEMENTS = {
    "contexto latinoamericano": "entorno regional",
    "innovacion": "desarrollo creativo",
    "innovación": "desarrollo creativo",
    "innovaci\ufffdn": "desarrollo creativo",
    "inmersion": "experiencia de juego",
    "inmersión": "experiencia de juego",
    "inmersi\ufffdn": "experiencia de juego",
    "narrativa interactiva": "diseño narrativo",
    "narrativas interactivas": "diseño narrativo",
    "accesibilidad": "diseño inclusivo",
    "sostenibilidad": "viabilidad del proyecto",
}

ACADEMIC_FILLER_PATTERNS = [
    r"\bes importante destacar que\b",
    r"\bcabe resaltar que\b",
    r"\ben este sentido,?\b",
    r"\bdesde esta perspectiva,?\b",
    r"\bresulta fundamental comprender que\b",
    r"\bde manera significativa\b",
]


SECTION_PLAN = [
    {
        "key": "introduccion",
        "title": "INTRODUCCIÓN",
        "target_words": "1.200 a 1.600 palabras",
        "min_words": 1000,
        "instruction": "Redacta solo la introducción. Debe abrir con una escena o pregunta concreta, contextualizar el tema y conectar con la asignatura, la competencia y el resultado de aprendizaje.",
    },
    {
        "key": "ejes_1",
        "title": "EJES ARTICULADORES",
        "target_words": "1.800 a 2.300 palabras",
        "min_words": 1500,
        "instruction": "Redacta la primera mitad de los ejes articuladores. Enfócate en fundamentos conceptuales, relevancia formativa, autores y relación con el contexto profesional.",
    },
    {
        "key": "ejes_2",
        "title": "EJES ARTICULADORES",
        "target_words": "1.800 a 2.300 palabras",
        "min_words": 1500,
        "instruction": "Redacta la segunda mitad de los ejes articuladores. Continúa sin repetir lo anterior; profundiza en aplicaciones, dilemas, prácticas académicas y ejemplos situados.",
    },
    {
        "key": "ensayo_1",
        "title": "ENSAYOS DE PROFUNDIZACIÓN",
        "target_words": "1.300 a 1.700 palabras",
        "min_words": 1100,
        "instruction": "Redacta el primer ensayo de profundización con subtítulo propio. Analiza un problema aplicado del tema y su impacto en la formación académica.",
    },
    {
        "key": "ensayo_2",
        "title": "ENSAYOS DE PROFUNDIZACIÓN",
        "target_words": "1.300 a 1.700 palabras",
        "min_words": 1100,
        "instruction": "Redacta el segundo ensayo de profundización con subtítulo propio. Analiza una tensión del tema en contextos laborales, digitales o institucionales.",
    },
    {
        "key": "ensayo_3",
        "title": "ENSAYOS DE PROFUNDIZACIÓN",
        "target_words": "1.300 a 1.700 palabras",
        "min_words": 1100,
        "instruction": "Redacta el tercer ensayo de profundización con subtítulo propio. Conecta el tema con ética, ciudadanía, responsabilidad social o proyecto de vida.",
    },
    {
        "key": "conclusiones",
        "title": "CONCLUSIONES",
        "target_words": "900 a 1.200 palabras",
        "min_words": 750,
        "instruction": "Redacta solo las conclusiones. Sintetiza aprendizajes, implicaciones profesionales y cierre formativo sin frases genéricas.",
    },
    {
        "key": "bibliografia",
        "title": "BIBLIOGRAFÍA",
        "target_words": "20 a 30 referencias posteriores a 2020",
        "min_words": 250,
        "instruction": "Redacta únicamente la bibliografía en formato APA 7. Incluye entre 20 y 30 referencias reales o ampliamente reconocibles, todas publicadas de 2021 en adelante. No incluyas referencias de 2020 ni de años anteriores. No inventes DOI ni URL.",
    },
]


SECTION_PLANS_BY_LEVEL = {
    "pregrado": SECTION_PLAN,
    "especializacion": [
        {
            "key": "introduccion",
            "title": "INTRODUCCION",
            "target_words": "1.800 a 2.500 palabras",
            "min_words": 1500,
            "instruction": "Redacta solo la introduccion. Plantea el problema desde una perspectiva de posgrado profesional, con rigor conceptual y conexion con la practica avanzada.",
        },
        {
            "key": "capitulos_1",
            "title": "CAPITULOS TEORICO-APLICADOS",
            "target_words": "2.500 a 3.500 palabras",
            "min_words": 2000,
            "instruction": "Redacta la primera mitad de los capitulos teorico-aplicados. Desarrolla fundamentos, autores, categorias y relacion con problemas profesionales avanzados.",
        },
        {
            "key": "capitulos_2",
            "title": "CAPITULOS TEORICO-APLICADOS",
            "target_words": "2.500 a 3.500 palabras",
            "min_words": 2000,
            "instruction": "Redacta la segunda mitad de los capitulos teorico-aplicados. Profundiza en aplicacion, tensiones, criterios de decision y contexto colombiano o regional.",
        },
        {
            "key": "analisis_1",
            "title": "ANALISIS CRITICOS",
            "target_words": "1.500 a 2.000 palabras",
            "min_words": 1200,
            "instruction": "Redacta el primer analisis critico con una tesis clara, soporte bibliografico y lectura profesional avanzada.",
        },
        {
            "key": "analisis_2",
            "title": "ANALISIS CRITICOS",
            "target_words": "1.500 a 2.000 palabras",
            "min_words": 1200,
            "instruction": "Redacta el segundo analisis critico. Contrasta enfoques, identifica implicaciones profesionales y evita repetir el analisis anterior.",
        },
        {
            "key": "analisis_3",
            "title": "ANALISIS CRITICOS",
            "target_words": "1.500 a 2.000 palabras",
            "min_words": 1200,
            "instruction": "Redacta el tercer analisis critico. Integra una tension etica, institucional, normativa o metodologica propia del nivel de especializacion.",
        },
        {
            "key": "conclusiones",
            "title": "CONCLUSIONES",
            "target_words": "1.200 a 1.600 palabras",
            "min_words": 900,
            "instruction": "Redacta solo las conclusiones. Sintetiza aportes conceptuales, implicaciones profesionales avanzadas y cierre argumentativo.",
        },
        {
            "key": "bibliografia",
            "title": "BIBLIOGRAFIA",
            "target_words": "12 a 18 referencias verificables posteriores a 2020",
            "min_words": 350,
            "instruction": "Redacta unicamente la bibliografia en APA 7. Incluye entre 12 y 18 referencias reales, verificables y pertinentes, publicadas de 2021 en adelante. Prioriza bibliografia del silabo, documentacion oficial, organismos reconocidos y fuentes institucionales. No inventes autores, DOI, URL, volumenes ni paginas.",
        },
    ],
    "diplomado": [
        {
            "key": "introduccion",
            "title": "INTRODUCCION",
            "target_words": "1.200 a 1.800 palabras",
            "min_words": 1000,
            "instruction": "Redacta solo la introduccion. Contextualiza desde un problema profesional concreto, la utilidad practica y las competencias del diplomado.",
        },
        {
            "key": "modulos_1",
            "title": "MODULOS TEMATICOS",
            "target_words": "2.000 a 2.800 palabras",
            "min_words": 1600,
            "instruction": "Redacta los primeros modulos tematicos con orientacion a competencias, herramientas aplicables, criterios de decision y transferencia al contexto laboral.",
        },
        {
            "key": "modulos_2",
            "title": "MODULOS TEMATICOS",
            "target_words": "2.000 a 2.800 palabras",
            "min_words": 1600,
            "instruction": "Redacta los modulos tematicos restantes. Integra escenarios profesionales, errores frecuentes, protocolos o metodologias utiles.",
        },
        {
            "key": "casos",
            "title": "CASOS DE APLICACION PROFESIONAL",
            "target_words": "3.000 a 4.500 palabras",
            "min_words": 2400,
            "instruction": "Redacta tres casos de aplicacion profesional con escenario, analisis, ruta de intervencion, tension profesional y lecciones transferibles.",
        },
        {
            "key": "conclusiones",
            "title": "CONCLUSIONES",
            "target_words": "800 a 1.200 palabras",
            "min_words": 650,
            "instruction": "Redacta solo las conclusiones. Sintetiza competencias profesionales, orientaciones practicas y tendencias de actualizacion continua.",
        },
        {
            "key": "bibliografia",
            "title": "BIBLIOGRAFIA",
            "target_words": "15 a 25 referencias posteriores a 2020",
            "min_words": 220,
            "instruction": "Redacta unicamente la bibliografia en APA 7. Incluye entre 15 y 25 referencias reales, priorizando normas, guias y fuentes aplicadas recientes.",
        },
    ],
    "maestria": [
        {
            "key": "introduccion",
            "title": "INTRODUCCION",
            "target_words": "2.500 a 3.500 palabras",
            "min_words": 2000,
            "instruction": "Redacta solo la introduccion. Construye el tema como problema intelectual, con estado de la cuestion, posicionamiento teorico y hoja de ruta argumentativa.",
        },
        {
            "key": "capitulos_1",
            "title": "CAPITULOS DE DESARROLLO",
            "target_words": "3.000 a 4.000 palabras",
            "min_words": 2400,
            "instruction": "Redacta los primeros capitulos de desarrollo con estructura investigativa: problema, revision critica, analisis, implicaciones, condiciones de validez y sintesis.",
        },
        {
            "key": "capitulos_2",
            "title": "CAPITULOS DE DESARROLLO",
            "target_words": "3.000 a 4.000 palabras",
            "min_words": 2400,
            "instruction": "Redacta los capitulos de desarrollo restantes. Profundiza en debates, genealogia conceptual, evidencia, limitaciones y aporte al argumento general.",
        },
        {
            "key": "discusiones",
            "title": "DISCUSIONES TEORICAS",
            "target_words": "5.000 a 7.000 palabras",
            "min_words": 3500,
            "instruction": "Redacta tres discusiones teoricas con tesis propia: una teorica, una metodologica y una aplicada o de politica.",
        },
        {
            "key": "conclusiones",
            "title": "CONCLUSIONES",
            "target_words": "1.800 a 2.500 palabras",
            "min_words": 1400,
            "instruction": "Redacta solo las conclusiones. Integra contribuciones, limitaciones, agenda de investigacion y cierre intelectualmente denso.",
        },
        {
            "key": "bibliografia",
            "title": "BIBLIOGRAFIA",
            "target_words": "40 a 55 referencias posteriores a 2020",
            "min_words": 450,
            "instruction": "Redacta unicamente la bibliografia en APA 7. Incluye entre 40 y 55 referencias reales y pertinentes.",
        },
    ],
}

SECTION_PLANS_BY_LEVEL["curso_rapido"] = SECTION_PLAN
SECTION_PLANS_BY_LEVEL["curso_externos_profesional"] = SECTION_PLANS_BY_LEVEL["diplomado"]


@dataclass
class CoursePlan:
    asignatura: str
    programa: str
    escuela: str
    semestre: str
    ciclo: str
    modalidad: str
    creditos: str
    competencias: str
    resultados_aprendizaje: str
    pregunta_problema: str
    temas: List[str]
    bibliografia_silabo: List[str]
    nivel: str = "pregrado"


@dataclass
class CourseInfo:
    asignatura: str
    programa: str
    escuela: str
    semestre: str
    temas: List[str]


@dataclass
class SyllabusParseResult:
    program: str
    school: str
    modality: str
    coursesDetected: List[CourseInfo]
    selectedCourse: CourseInfo | None


FIELD_LABELS = {
    "CODIGO",
    "CODIGO CURSO",
    "NOMBRE",
    "ASIGNATURA",
    "CURSO",
    "SEMESTRE",
    "SEMESTRE NIVEL",
    "NIVEL",
    "TIPO",
    "COMPONENTE",
    "CALIFICACION",
    "MODALIDAD",
    "MODALIDAD DEL PROGRAMA",
    "CREDITOS",
    "CREDITOS ACADEMICOS",
    "PROGRAMA",
    "ESCUELA",
}

PROGRAM_REJECTS = {
    "PRESENCIAL",
    "VIRTUAL",
    "DISTANCIA",
    "DUAL",
    "HIBRIDO",
    "HÍBRIDO",
    "MODALIDAD DEL PROGRAMA",
    "NOMBRE",
    "PROGRAMA",
    "ESCUELA",
}

TOPIC_REJECT_PHRASES = [
    "PREGUNTA PROBLEMA",
    "ESTRATEGIA DIDACTICA",
    "DURACION",
    "TOTAL HORAS",
    "CREDITOS",
    "TALLERES PRACTICOS",
    "PREGUNTAS ORIENTADORAS",
    "IMPLEMENTACION DE BASES DE DATOS",
    "DESARROLLO DE SEMINARIO ALEMAN",
]


def detect_all_courses(path: Path) -> List[CourseInfo]:
    if path.suffix.lower() == ".pdf":
        text = extract_pdf_text(path)
        return _detect_courses_from_text(text, path)

    parsed = parse_syllabus_docx(path)
    if parsed.coursesDetected:
        return parsed.coursesDetected

    plan = extract_course_plan_from_text(extract_docx_text(path), path)
    return [CourseInfo(plan.asignatura, plan.programa, plan.escuela, plan.semestre, plan.temas)]


def parse_syllabus_docx(path: Path) -> SyllabusParseResult:
    doc = Document(path)
    rows = extract_ordered_docx_rows(doc)
    full_text = extract_docx_text(path)
    program = detect_global_program(rows)
    if not program:
        program = extract_program_fallback_from_flat_docx_text(full_text)
    school = detect_global_field(rows, "Escuela")
    modality = detect_global_modality(rows)

    block_ranges = detect_course_block_ranges(rows)
    if not block_ranges:
        block_ranges = [(0, len(rows))]

    courses: List[CourseInfo] = []
    for start, end in block_ranges:
        block_rows = rows[start:end]
        block_text = "\n".join(" | ".join(cell for cell in row if clean_text(cell)) for row in block_rows)
        subject = extract_value_from_block(block_rows, "Nombre", ["Nombre de la asignatura", "Nombre de la materia", "Asignatura", "Curso"])
        subject = sanitize_subject(subject)
        if not subject:
            continue

        semester = extract_value_from_block(block_rows, "Semestre", ["Semestre (Nivel)", "Nivel", "Semestre/Nivel"])
        topics = extract_topics_from_course_block(block_rows, block_text)
        if not topics:
            topics = extract_topics_from_text_block(block_text)

        course_program = program or detect_global_program(block_rows)
        course_school = school or detect_global_field(block_rows, "Escuela")
        courses.append(CourseInfo(
            asignatura=subject,
            programa=course_program,
            escuela=course_school,
            semestre=semester,
            temas=topics[:5],
        ))

    selected = next((course for course in courses if course.temas), courses[0] if courses else None)
    return SyllabusParseResult(
        program=program,
        school=school,
        modality=modality,
        coursesDetected=courses,
        selectedCourse=selected,
    )


def extract_ordered_docx_rows(doc: Document) -> List[List[str]]:
    ordered_rows: List[List[str]] = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, doc)
            text = clean_text(paragraph.text)
            if text:
                ordered_rows.append([text])
        elif child.tag.endswith("}tbl"):
            table = Table(child, doc)
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                if any(cells):
                    ordered_rows.append(cells)
    return ordered_rows


def detect_course_block_ranges(rows: List[List[str]]) -> List[tuple[int, int]]:
    starts = []
    for index, row in enumerate(rows):
        row_norm = normalize_heading(" ".join(row))
        if "1.1 IDENTIFICACION DEL CURSO" in row_norm or "IDENTIFICACION DEL CURSO" in row_norm:
            if starts and index == starts[-1] + 1:
                starts[-1] = index
            else:
                starts.append(index)

    ranges = []
    for position, start in enumerate(starts):
        end = starts[position + 1] if position + 1 < len(starts) else len(rows)
        ranges.append((start, end))
    return ranges


def _looks_like_institutional_school_row(text: str) -> bool:
    """Nombre de escuela/facultad en una sola celda (plantillas CUN sin etiqueta 'Escuela')."""
    if not clean_text(text):
        return False
    n = normalize_heading(text)
    markers = (
        "ESCUELA",
        "FACULTAD",
        "INSTITUTO",
        "CORPORACION",
        "CORPORACI",  # truncado por encoding
        "UNIVERSIDAD",
        "COLEGIO",
        "CENTRO DE",
    )
    return any(m in n for m in markers)


def _is_obviously_not_program_name(text: str) -> bool:
    """Valores sueltos de modalidad u opciones entre escuela y programa."""
    t = clean_text(text)
    if len(t) <= 3:
        return True
    n = normalize_heading(t)
    if n in {"X", "SI", "NO"}:
        return True
    if n in {"PRESENCIAL", "VIRTUAL", "DISTANCIA", "DUAL", "HIBRIDO", "HÍBRIDO"}:
        return True
    return False


def detect_program_from_school_program_stacked_rows(rows: List[List[str]]) -> str:
    """Dos filas consecutivas de una sola celda: institución (escuela/…) y nombre del programa."""
    for i in range(len(rows) - 1):
        r0, r1 = rows[i], rows[i + 1]
        if len(r0) != 1 or len(r1) != 1:
            continue
        a = clean_text(r0[0])
        b = clean_text(r1[0])
        if not a or not b:
            continue
        if not _looks_like_institutional_school_row(a):
            continue
        if _is_obviously_not_program_name(b):
            continue
        if is_valid_program_value(b):
            return b
    return ""


def detect_global_program(rows: List[List[str]]) -> str:
    for row_index, row in enumerate(rows):
        for cell_index, cell in enumerate(row):
            if normalize_label_token(cell) == "PROGRAMA":
                candidates = collect_following_values(rows, row_index, cell_index)
                for candidate in candidates:
                    cleaned = clean_text(candidate)
                    if is_valid_program_value(cleaned):
                        return cleaned
    stacked = detect_program_from_school_program_stacked_rows(rows)
    if stacked:
        return stacked
    return ""


def detect_global_field(rows: List[List[str]], label: str) -> str:
    label_norm = normalize_label_token(label)
    for row_index, row in enumerate(rows):
        for cell_index, cell in enumerate(row):
            if normalize_label_token(cell) == label_norm:
                for candidate in collect_following_values(rows, row_index, cell_index):
                    candidate = clean_text(candidate)
                    if candidate and normalize_heading(candidate) not in FIELD_LABELS:
                        return candidate
    return ""


def detect_global_modality(rows: List[List[str]]) -> str:
    for row in rows:
        normalized = [normalize_heading(cell) for cell in row]
        if any("MODALIDAD DEL PROGRAMA" in cell or cell == "MODALIDAD" for cell in normalized):
            for cell in row:
                if normalize_heading(cell) in {"PRESENCIAL", "VIRTUAL", "DISTANCIA", "DUAL", "HIBRIDO"}:
                    return clean_text(cell)
    return ""


def collect_following_values(rows: List[List[str]], row_index: int, cell_index: int) -> List[str]:
    values: List[str] = []
    row = rows[row_index]
    for value in row[cell_index + 1:]:
        if clean_text(value):
            values.append(value)
    if row_index + 1 < len(rows):
        for value in rows[row_index + 1]:
            if clean_text(value):
                values.append(value)
    return values


def is_valid_program_value(value: str) -> bool:
    cleaned = clean_text(value)
    normalized = normalize_heading(cleaned)
    if not cleaned or normalized in PROGRAM_REJECTS or normalized in FIELD_LABELS:
        return False
    if any(word in normalized for word in ["ESPECIALIZACION", "PREGRADO", "MAESTRIA", "DIPLOMADO"]):
        return True
    meaningful_words = [word for word in re.findall(r"[A-ZÁÉÍÓÚÑa-záéíóúñ]+", cleaned) if len(word) > 2]
    # Nombres cortos pero válidos (ej. "Técnica Profesional X") o títulos largos
    if len(meaningful_words) >= 2 and len(cleaned) >= 12:
        return True
    return len(meaningful_words) > 3


def extract_program_fallback_from_flat_docx_text(text: str) -> str:
    """Último recurso: celdas 'Programa | valor' o dos líneas consecutivas escuela → programa (TABLAs en extract_docx_text)."""
    lines = text.splitlines()
    for raw in lines:
        line = clean_text(raw)
        if not line or line.upper().startswith("TABLA"):
            continue
        cells = [clean_text(c) for c in re.split(r"\s*\|\s*", line) if clean_text(c)]
        for i, cell in enumerate(cells):
            if normalize_label_token(cell) == "PROGRAMA" and i + 1 < len(cells):
                val = clean_text(cells[i + 1])
                if is_valid_program_value(val):
                    return val
    for i, raw in enumerate(lines[:-1]):
        line, nxt = clean_text(raw), clean_text(lines[i + 1])
        if not line or not nxt or "|" in line or "|" in nxt:
            continue
        if line.upper().startswith("TABLA") or lines[i + 1].upper().startswith("TABLA"):
            continue
        if _looks_like_institutional_school_row(line) and is_valid_program_value(nxt) and not _is_obviously_not_program_name(nxt):
            return nxt
    return ""


def extract_value_from_block(rows: List[List[str]], label: str, aliases: List[str] | None = None) -> str:
    labels = {normalize_label_token(label), *(normalize_label_token(alias) for alias in (aliases or []))}
    for row_index, row in enumerate(rows):
        for cell_index, cell in enumerate(row):
            cell_norm = normalize_label_token(cell)
            if cell_norm in labels:
                candidates = collect_label_value_fragments(rows, row_index, cell_index)
                value = join_value_fragments(candidates)
                if value:
                    return value
    return ""


def collect_label_value_fragments(rows: List[List[str]], row_index: int, cell_index: int) -> List[str]:
    fragments: List[str] = []
    row = rows[row_index]
    for value in row[cell_index + 1:]:
        cleaned = clean_text(value)
        if not cleaned:
            continue
        if normalize_heading(cleaned) in FIELD_LABELS:
            break
        fragments.append(cleaned)
    if fragments and not (len(fragments) == 1 and len(fragments[0]) == 1):
        return fragments

    for next_row in rows[row_index + 1: row_index + 3]:
        for value in next_row:
            cleaned = clean_text(value)
            if not cleaned:
                continue
            if normalize_heading(cleaned) in FIELD_LABELS:
                if fragments:
                    return fragments
                continue
            fragments.append(cleaned)
        if fragments:
            break
    return fragments


def join_value_fragments(fragments: List[str]) -> str:
    useful = []
    seen = set()
    for fragment in fragments:
        if normalize_heading(fragment) in FIELD_LABELS:
            continue
        key = normalize_heading(fragment)
        if key in seen:
            continue
        useful.append(fragment)
        seen.add(key)
    if not useful:
        return ""
    result = useful[0]
    for fragment in useful[1:]:
        if len(result) == 1 and fragment[:1].islower():
            result += fragment
        elif len(result) == 1 and fragment[:1].isalpha():
            result += fragment
        else:
            result += " " + fragment
    return clean_text(result)


def sanitize_subject(value: str) -> str:
    cleaned = clean_text(value)
    normalized = normalize_heading(cleaned)
    if not cleaned or normalized in FIELD_LABELS or normalized in PROGRAM_REJECTS:
        return ""
    if len(cleaned) == 1:
        return ""
    return cleaned


def normalize_label_token(value: str) -> str:
    return re.sub(r"[^A-Z0-9]+", " ", normalize_heading(value)).strip()


def extract_topics_from_course_block(rows: List[List[str]], block_text: str) -> List[str]:
    start = None
    for index, row in enumerate(rows):
        if "ESTRUCTURA TEMATICA" in normalize_heading(" ".join(row)):
            start = index
            break
    if start is None:
        return []

    content_index = None
    topics: List[str] = []
    for row in rows[start + 1:]:
        row_norm = normalize_heading(" ".join(row))
        if "ESTRATEGIA DIDACTICA" in row_norm and "CONTENIDOS" not in row_norm:
            break
        if "BIBLIOGRAFIA" in row_norm or "OBSERVACIONES" in row_norm:
            break

        normalized_cells = [normalize_heading(cell) for cell in row]
        if any(cell == "CONTENIDOS" for cell in normalized_cells):
            content_index = normalized_cells.index("CONTENIDOS")
            continue

        candidate = ""
        if content_index is not None and content_index < len(row):
            candidate = clean_text(row[content_index])
        elif len(row) >= 2:
            candidate = clean_text(row[1])
        if is_valid_topic(candidate) and candidate not in topics:
            topics.append(candidate)
        if len(topics) == 5:
            return topics

    return topics[:5] if topics else extract_topics_from_text_block(block_text)


def extract_topics_from_text_block(text: str) -> List[str]:
    topics: List[str] = []
    lines = [clean_text(line) for line in text.splitlines() if clean_text(line)]
    in_structure = False
    content_mode = False
    for line in lines:
        normalized = normalize_heading(line)
        if "ESTRUCTURA TEMATICA" in normalized:
            in_structure = True
            continue
        if in_structure and "ESTRATEGIA DIDACTICA" in normalized and "CONTENIDOS" not in normalized:
            break
        if not in_structure:
            continue
        if "CONTENIDOS" in normalized:
            content_mode = True
            continue
        if content_mode:
            for candidate in split_topic_candidates(line):
                candidate = clean_text(candidate)
                if is_valid_topic(candidate) and candidate not in topics:
                    topics.append(candidate)
                if len(topics) == 5:
                    return topics
    return topics[:5]


def is_valid_topic(value: str) -> bool:
    cleaned = clean_text(value)
    normalized = normalize_heading(cleaned)
    if not cleaned or normalized in FIELD_LABELS:
        return False
    if any(phrase in normalized for phrase in TOPIC_REJECT_PHRASES):
        return False
    if cleaned.startswith("¿") or re.fullmatch(r"\d+", cleaned):
        return False
    if len(cleaned.split()) < 2 or len(cleaned.split()) > 14:
        return False
    return True


def _detect_courses_from_text(text: str, path: Path) -> List[CourseInfo]:
    plan = extract_course_plan_from_text(text, path)
    return [CourseInfo(
        asignatura=plan.asignatura,
        programa=plan.programa,
        escuela=plan.escuela,
        semestre=plan.semestre,
        temas=plan.temas,
    )]


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[áàäâ]", "a", value)
    value = re.sub(r"[éèëê]", "e", value)
    value = re.sub(r"[íìïî]", "i", value)
    value = re.sub(r"[óòöô]", "o", value)
    value = re.sub(r"[úùüû]", "u", value)
    value = re.sub(r"ñ", "n", value)
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "documento"


def normalize_for_matching(value: str) -> str:
    value = unicodedata.normalize("NFKD", value or "")
    value = "".join(char for char in value if not unicodedata.combining(char))
    value = value.lower()
    value = value.replace("Ã¡", "a").replace("Ã©", "e").replace("Ã­", "i").replace("Ã³", "o").replace("Ãº", "u")
    value = value.replace("Ã±", "n").replace("Ã", "i").replace("Ã‰", "e").replace("Ã“", "o")
    return re.sub(r"\s+", " ", value)


def detect_academic_level(plan: CoursePlan, syllabus_text: str, requested_level: str) -> str:
    if requested_level != "auto":
        return requested_level

    haystack = normalize_for_matching(
        " ".join(
            [
                plan.programa,
                plan.asignatura,
                plan.escuela,
                plan.semestre,
                syllabus_text[:12000],
            ]
        )
    )
    checks = [
        ("maestria", ["maestria", "magister", "master"]),
        ("especializacion", ["especializacion", "especialista", "posgrado", "postgrado"]),
        ("diplomado", ["diplomado", "educacion continua", "formacion continua"]),
        ("pregrado", ["pregrado", "profesional universitario", "universitario", "semestre"]),
    ]
    for level, markers in checks:
        if any(marker in haystack for marker in markers):
            return level
    return "pregrado"


def resolve_prompt_path(level: str, prompt_override: str = "") -> Path:
    if prompt_override:
        return Path(prompt_override)
    return PROMPT_PATHS[level]


def get_section_plan(level: str) -> List[dict]:
    return SECTION_PLANS_BY_LEVEL[level]


def iter_table_rows(doc: Document) -> Iterable[List[str]]:
    for table in doc.tables:
        for row in table.rows:
            yield [clean_text(cell.text) for cell in row.cells]


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"TABLA {table_index}")
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = "\n".join(clean_text(line) for line in text.splitlines() if clean_text(line))
        if text:
            parts.append(f"PAGINA {index}\n{text}")
    return "\n\n".join(parts)


def extract_input_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    raise ValueError("Formato no soportado. Usa un archivo .docx o .pdf")


def extract_course_plan(path: Path, subject_override: str = "", semester_override: str = "", topics_override: str = "") -> CoursePlan:
    if path.suffix.lower() == ".pdf":
        return extract_course_plan_from_text(extract_pdf_text(path), path, subject_override, semester_override, topics_override)

    doc = Document(path)
    rows = list(iter_table_rows(doc))
    if not rows:
        return extract_course_plan_from_text(extract_docx_text(path), path, subject_override, semester_override, topics_override)

    def cell_contains(cell_text: str, label: str) -> bool:
        cell_norm = normalize_for_matching(cell_text)
        label_norm = normalize_for_matching(label)
        return label_norm in cell_norm

    def cell_matches_program_label_cell(cell_text: str, label: str) -> bool:
        """Celda de etiqueta Programa (no 'Modalidad del programa' ni 'Tipo de programa')."""
        head_tokens = normalize_label_token(cell_text).split()
        need_tokens = normalize_label_token(label).split()
        if not head_tokens or not need_tokens or need_tokens[0] != "PROGRAMA":
            cell_norm = normalize_for_matching(cell_text)
            lbl = normalize_for_matching(label)
            return bool(cell_norm and lbl and lbl in cell_norm)
        if head_tokens[0] != "PROGRAMA":
            return False
        if "MODALIDAD" in head_tokens:
            return False
        for i, tok in enumerate(need_tokens):
            if i >= len(head_tokens) or head_tokens[i] != tok:
                return False
        return True

    def value_after_flexible(label: str, aliases: List[str] | None = None) -> str:
        all_labels = [label] + (aliases or [])
        for row in rows:
            for index, cell in enumerate(row):
                if any(cell_contains(cell, lbl) for lbl in all_labels):
                    if index + 1 < len(row):
                        candidate = clean_text(row[index + 1])
                        if candidate and not any(cell_contains(candidate, lbl) for lbl in all_labels):
                            return candidate
                    if row_index_of_cell(row, cell) + 1 < len(rows):
                        next_row = rows[row_index_of_cell(row, cell) + 1]
                        if next_row:
                            candidate = clean_text(next_row[0])
                            if candidate and len(candidate) > 1:
                                return candidate
        return ""

    def row_index_of_cell(row: List[str], cell: str) -> int:
        for i, r in enumerate(rows):
            if r is row:
                return i
        return -1

    def find_in_row_pattern(label: str, aliases: List[str] | None = None) -> str:
        all_labels = [label] + (aliases or [])
        matcher = cell_matches_program_label_cell if normalize_for_matching(label) == "programa" else cell_contains
        for row in rows:
            for index, cell in enumerate(row):
                if any(matcher(cell, lbl) for lbl in all_labels):
                    if index + 1 < len(row):
                        val = clean_text(row[index + 1])
                        if val and len(val) > 1:
                            return val
        return ""

    escuela = find_in_row_pattern("Escuela")
    programa = find_in_row_pattern("Programa", ["Programa académico", "Programa de"])
    if not programa:
        programa = detect_program_from_school_program_stacked_rows(rows)
    asignatura = subject_override or find_in_row_pattern("Nombre", ["Nombre de la asignatura", "Nombre de la materia", "Asignatura", "Curso"])
    semestre = semester_override or find_in_row_pattern("Semestre", ["Semestre (Nivel)", "Nivel", "Semestre/Nivel"])
    modalidad = find_in_row_pattern("Modalidad", ["Modalidad del programa"])
    creditos = ""
    competencias = ""
    resultados = ""
    pregunta = ""
    temas: List[str] = []

    for row_index, row in enumerate(rows):
        lowered = [normalize_for_matching(cell) for cell in row]
        row_text = " ".join(row).lower()

        if row and any(cell_contains(row[0], "modalidad") for _ in [1]) and len(row) >= 3:
            for ci, rc in enumerate(row):
                if "modalidad" in normalize_for_matching(rc) and ci + 1 < len(row):
                    if row[ci + 1].strip().upper() == "X" and ci + 2 < len(row):
                        modalidad = row[ci + 1] if len(row[ci + 1]) > len(row[ci + 2]) else row[ci + 2]
                    elif row[ci + 1].strip().upper() == "X":
                        modalidad = row[ci] if ci > 0 else ""

        if any("competencia" in c for c in lowered) and any("resultado" in c for c in lowered) and row_index + 1 < len(rows):
            next_row = rows[row_index + 1]
            if len(next_row) >= 2:
                competencias = next_row[0]
                resultados = next_row[1]
            continue

        if len(row) >= 2 and row[0].strip().startswith("¿") and row[1] and normalize_for_matching(row[1]) not in {"contenidos", "total horas", ""}:
            pregunta = pregunta or row[0]
            tema = clean_text(row[1])
            if tema and tema not in temas:
                temas.append(tema)

        if len(row) >= 3:
            last_cell = clean_text(row[-1])
            if re.fullmatch(r"\d+", last_cell) and any("credito" in c for c in lowered):
                creditos = last_cell

    paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    bibliography = extract_bibliography(paragraphs)

    parsed = parse_syllabus_docx(path)
    if parsed.selectedCourse:
        asignatura = subject_override or parsed.selectedCourse.asignatura or asignatura
        programa = parsed.program or parsed.selectedCourse.programa or programa
        escuela = parsed.school or parsed.selectedCourse.escuela or escuela
        semestre = semester_override or parsed.selectedCourse.semestre or semestre
        if not topics_override and parsed.selectedCourse.temas:
            temas = parsed.selectedCourse.temas[:5]

    if not asignatura:
        asignatura = path.stem
    if topics_override:
        temas = parse_topics_override(topics_override)
    if not temas:
        temas = extract_topics_from_table_rows(rows)
    if not temas:
        temas = infer_topics_from_text("\n".join(paragraphs))
    if len(temas) > 5:
        temas = temas[:5]

    return CoursePlan(
        asignatura=asignatura,
        programa=programa,
        escuela=escuela,
        semestre=semestre,
        ciclo="Técnico" if "TÉCNICA" in programa.upper() else "",
        modalidad=modalidad,
        creditos=creditos,
        competencias=competencias,
        resultados_aprendizaje=resultados,
        pregunta_problema=pregunta,
        temas=temas,
        bibliografia_silabo=bibliography,
    )


def extract_course_plan_from_text(text: str, path: Path, subject_override: str = "", semester_override: str = "", topics_override: str = "") -> CoursePlan:
    lines = [clean_text(line) for line in text.splitlines() if clean_text(line)]

    def find_after(labels: List[str]) -> str:
        for line in lines:
            for label in labels:
                pattern = rf"{re.escape(label)}\s*:?\s*(.+)$"
                match = re.search(pattern, line, flags=re.IGNORECASE)
                if match:
                    return clean_text(match.group(1))
        return ""

    escuela = find_after(["Escuela"])
    programa = find_after(["Programa"])
    asignatura = subject_override or find_after(["Nombre", "Asignatura", "Curso"])
    semestre = semester_override or find_after(["Semestre", "Nivel"])
    modalidad = find_after(["Modalidad"])
    creditos = find_after(["Creditos", "Créditos"])
    competencias = extract_section_from_text(text, ["COMPETENCIAS", "ELEMENTOS DE LA MACROCOMPETENCIA"], ["ESTRUCTURA TEMATICA", "ESTRUCTURA TEMÁTICA", "RESULTADOS"])
    resultados = extract_section_from_text(text, ["RESULTADOS DE APRENDIZAJE", "RESULTADO DE APRENDIZAJE"], ["ESTRUCTURA TEMATICA", "ESTRUCTURA TEMÁTICA", "JUSTIFICACION", "JUSTIFICACIÓN"])
    pregunta = extract_problem_question(text)
    temas = parse_topics_override(topics_override) if topics_override else extract_topics_from_text(text)
    bibliography = extract_bibliography(lines)

    if not asignatura:
        asignatura = path.stem
    if len(temas) > 5:
        temas = temas[:5]

    return CoursePlan(
        asignatura=asignatura,
        programa=programa,
        escuela=escuela,
        semestre=semestre,
        ciclo="Técnico" if "TECNICA" in normalize_heading(programa) else "",
        modalidad=modalidad,
        creditos=creditos,
        competencias=competencias,
        resultados_aprendizaje=resultados,
        pregunta_problema=pregunta,
        temas=temas,
        bibliografia_silabo=bibliography,
    )


def extract_topics_from_table_rows(rows: List[List[str]]) -> List[str]:
    topics: List[str] = []
    in_structure = False
    stop_markers = {
        "ESTRATEGIA DIDACTICA",
        "ESTRATEGIA DIDACTICA",
        "MECANISMOS Y ESTRATEGIAS DE EVALUACION",
        "MECANISMOS Y ESTRATEGIAS DE EVALUACION",
        "BIBLIOGRAFIA",
        "BIBLIOGRAFIA",
        "OBSERVACIONES",
        "OBSERVACIONES",
    }

    for row in rows:
        if not row:
            continue
        first_cell = normalize_heading(row[0])
        if "ESTRUCTURA TEMATICA" in first_cell or "ESTRUCTURA TEMATICA" in first_cell:
            in_structure = True
            continue
        if in_structure and first_cell in stop_markers:
            break
        if not in_structure:
            continue

        for cell in row[1:]:
            cleaned = clean_text(cell)
            normalized = normalize_heading(cleaned)
            if (
                cleaned
                and 2 <= len(cleaned.split()) <= 12
                and normalized not in {"CONTENIDOS", "TOTAL HORAS", "PREGUNTA PROBLEMA", ""}
                and not cleaned.startswith("¿")
                and not re.fullmatch(r"\d+", cleaned)
                and cleaned not in topics
            ):
                topics.append(cleaned)
            if len(topics) == 5:
                return topics

    return topics[:5]


def extract_bibliography(paragraphs: List[str]) -> List[str]:
    items: List[str] = []
    in_biblio = False
    for text in paragraphs:
        upper = text.upper()
        if upper == "BIBLIOGRAFÍA" or upper == "BIBLIOGRAFIA":
            in_biblio = True
            continue
        if in_biblio and upper == "OBSERVACIONES":
            break
        if in_biblio and len(text) > 20:
            items.append(text)
    return items


def parse_topics_override(value: str) -> List[str]:
    return [clean_text(part) for part in re.split(r"[|;]", value) if clean_text(part)][:5]


def extract_section_from_text(text: str, start_markers: List[str], end_markers: List[str]) -> str:
    normalized_lines = [(line, normalize_heading(line)) for line in text.splitlines() if clean_text(line)]
    collecting = False
    parts = []
    for original, normalized in normalized_lines:
        if not collecting and any(marker in normalized for marker in start_markers):
            collecting = True
            continue
        if collecting and any(marker in normalized for marker in end_markers):
            break
        if collecting:
            parts.append(clean_text(original))
    return clean_text(" ".join(parts))


def extract_problem_question(text: str) -> str:
    match = re.search(r"(¿[^?]+\?)", text)
    return clean_text(match.group(1)) if match else ""


def extract_topics_from_text(text: str) -> List[str]:
    lines = [clean_text(line) for line in text.splitlines() if clean_text(line)]
    topics = []
    in_structure = False
    stop_markers = {
        "ESTRATEGIA DIDACTICA",
        "ESTRATEGIA DIDÁCTICA",
        "MECANISMOS Y ESTRATEGIAS DE EVALUACION",
        "MECANISMOS Y ESTRATEGIAS DE EVALUACIÓN",
        "BIBLIOGRAFIA",
        "BIBLIOGRAFÍA",
    }

    for line in lines:
        normalized = normalize_heading(line)
        if "ESTRUCTURA TEMATICA" in normalized:
            in_structure = True
            continue
        if in_structure and normalized in stop_markers:
            break
        if not in_structure:
            continue

        candidates = split_topic_candidates(line)
        for candidate in candidates:
            normalized_candidate = normalize_heading(candidate)
            if (
                2 <= len(candidate.split()) <= 8
                and normalized_candidate not in {"PREGUNTA PROBLEMA", "CONTENIDOS", "ESTRATEGIA DIDACTICA", "DURACION TCD TAE", "TOTAL HORAS"}
                and not candidate.startswith("¿")
                and not re.fullmatch(r"\d+", candidate)
                and candidate not in topics
            ):
                topics.append(candidate)
            if len(topics) == 5:
                return topics

    if len(topics) < 5:
        topics.extend(topic for topic in infer_topics_from_text(text) if topic not in topics)
    return topics[:5]


def split_topic_candidates(line: str) -> List[str]:
    if "|" in line:
        return [clean_text(part) for part in line.split("|") if clean_text(part)]
    numbered = re.split(r"(?:^|\s)(?:\d+[\.\)]|[a-eA-E][\.\)])\s+", line)
    return [clean_text(part) for part in numbered if clean_text(part)]


def infer_topics_from_text(text: str) -> List[str]:
    candidates = [
        "Escucha activa",
        "Comunicación verbal",
        "Comunicación no verbal",
        "Empatía",
        "Asertividad",
    ]
    return [topic for topic in candidates if topic.lower() in text.lower()][:5]


def load_system_prompt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def build_user_prompt(plan: CoursePlan, syllabus_text: str, topic: str, topic_index: int, total_topics: int) -> str:
    bibliography = format_recent_bibliography(plan.bibliografia_silabo)
    return f"""
Genera el documento temático {topic_index} de {total_topics}.

Tema central:
{topic}

Metadatos del curso:
{json.dumps(asdict(plan), ensure_ascii=False, indent=2)}

Bibliografía detectada en el sílabo:
{bibliography}

Sílabo completo extraído:
{syllabus_text[:30000]}

Instrucciones específicas:
- El documento debe estar centrado solo en el tema "{topic}", sin convertirse en un resumen general de toda la asignatura.
- Conserva la estructura: encabezado, INTRODUCCIÓN, EJES ARTICULADORES, ENSAYOS DE PROFUNDIZACIÓN, CONCLUSIONES, BIBLIOGRAFÍA.
- Extensión objetivo: 7.500 a 10.000 palabras. Si el límite de salida impide llegar a esa extensión, prioriza profundidad, coherencia y una bibliografía completa.
- Incluye solo referencias APA 7 reales y verificables, posteriores a 2020. No inventes fuentes, autores, DOI ni URL.
- Evita redundancia artificial: cada parrafo debe aportar una idea, criterio, comparacion, tension o implicacion nueva.
""".strip()


def build_section_prompt(
    plan: CoursePlan,
    syllabus_text: str,
    topic: str,
    topic_index: int,
    total_topics: int,
    section: dict,
    previous_context: str,
) -> str:
    bibliography = format_recent_bibliography(plan.bibliografia_silabo)
    previous_note = previous_context[-2500:] if previous_context else "Aun no hay secciones redactadas."
    return f"""
Vas a redactar una parte de un documento academico largo de nivel {plan.nivel}. Usa el prompt del sistema correspondiente a ese nivel academico como regla principal de estilo, profundidad, estructura y bibliografia.

Documento tematico {topic_index} de {total_topics}
Tema central: {topic}
Nivel academico seleccionado: {plan.nivel}
Seccion a redactar ahora: {section["title"]}
Extension obligatoria de esta respuesta: {section["target_words"]}

Instruccion de esta seccion:
{section["instruction"]}

Metadatos del curso:
{json.dumps(asdict(plan), ensure_ascii=False, indent=2)}

Bibliografia detectada en el silabo:
{bibliography}

Fragmento de contexto ya redactado para evitar repeticiones:
{previous_note}

Silabo completo extraido:
{syllabus_text[:30000]}

Reglas de salida para esta llamada:
- Escribe solo el contenido de la seccion solicitada.
- Respeta la estructura y el tono propios del nivel academico seleccionado.
- No repitas el encabezado institucional.
- No escribas el titulo principal de la seccion; el sistema lo agrega automaticamente.
- No escribas explicaciones sobre el proceso.
- No uses markdown, asteriscos, emojis ni listas esquematicas salvo en BIBLIOGRAFIA.
- Si la seccion es BIBLIOGRAFIA, escribe solo referencias APA 7, una por parrafo, y todas deben ser de 2021 en adelante.
- Si la seccion es BIBLIOGRAFIA, no incluyas referencias dudosas, sinteticas o inventadas. Si no conoces un DOI, URL, volumen, numero o paginas, omite ese dato.
- Si la seccion no es BIBLIOGRAFIA, integra citas narrativas o parenteticas solo cuando sean necesarias y verificables; si no hay seguridad bibliografica, usa formulaciones generales como "investigaciones contemporaneas" o "literatura especializada" sin inventar autores.
- No uses citas en texto de 2020 o anteriores aunque aparezcan en el silabo; pueden orientar el contexto, pero no deben figurar como referencia academica vigente.
- No cierres el documento si aun no estas en CONCLUSIONES.

Reglas editoriales anti-redundancia:
- No reformules la misma idea multiples veces ni uses parrafos con la misma estructura argumental.
- Cada parrafo debe aportar informacion nueva: concepto, criterio tecnico, comparacion, implicacion profesional, tension conceptual, caso aplicado o decision metodologica.
- Prioriza densidad conceptual sobre longitud artificial: no expandas definiciones si no agregas analisis real.
- Evita frases academicas genericas y muletillas como "es importante destacar", "en este sentido", "desde esta perspectiva" o equivalentes repetidos.
- No repitas constantemente las expresiones "contexto latinoamericano", "innovacion", "inmersion", "narrativa interactiva", "accesibilidad" ni "sostenibilidad". Usa esos conceptos solo cuando agreguen precision analitica.
- No cierres parrafos con conclusiones vacias; cierra con una consecuencia tecnica, pedagogica o profesional concreta.
""".strip()


def generate_document(client: OpenAI, model: str, system_prompt: str, user_prompt: str, max_tokens: int, temperature: float) -> str:
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    choice = response.choices[0]
    content = choice.message.content
    if content is None:
        finish = getattr(choice, "finish_reason", None)
        refusal = getattr(choice.message, "refusal", None)
        raise RuntimeError(
            "OpenAI devolvió message.content vacío (None). "
            f"finish_reason={finish!r}; refusal={refusal!r}. "
            "Revisa modelo, max_tokens, moderación o filtrado de contenido."
        )
    return content.strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\wÁÉÍÓÚáéíóúÑñÜü]+\b", text))


def reference_count(text: str) -> int:
    return len([line for line in text.splitlines() if re.search(r"\(\d{4}\)", line)])


def reference_years(text: str) -> List[int]:
    return [int(year) for year in re.findall(r"\((\d{4})[a-z]?\)", text)]


def recent_reference_count(text: str) -> int:
    return len([year for year in reference_years(text) if year > 2020])


def has_old_references(text: str) -> bool:
    years = reference_years(text)
    return any(year <= 2020 for year in years)


def normalize_for_similarity(text: str) -> str:
    value = unicodedata.normalize("NFKD", text.lower())
    value = "".join(char for char in value if not unicodedata.combining(char))
    value = re.sub(r"[^a-z0-9ñáéíóúü\s]", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def token_set(text: str) -> set[str]:
    stopwords = {
        "para", "como", "desde", "esta", "este", "estos", "estas", "porque", "sobre",
        "entre", "donde", "cuando", "tambien", "también", "puede", "debe", "hacia",
        "cada", "forma", "manera", "proceso", "tema", "nivel", "contexto", "analisis",
        "análisis", "profesional", "academico", "académico", "estudiante", "documento",
    }
    return {word for word in normalize_for_similarity(text).split() if len(word) > 4 and word not in stopwords}


def paragraph_similarity(left: str, right: str) -> float:
    left_tokens = token_set(left)
    right_tokens = token_set(right)
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / len(left_tokens | right_tokens)


def is_near_duplicate(paragraph: str, previous_paragraphs: List[str]) -> bool:
    normalized = normalize_for_similarity(paragraph)
    if len(normalized.split()) < 35:
        return False
    for previous in previous_paragraphs[-18:]:
        if paragraph_similarity(paragraph, previous) >= 0.82:
            return True
    return False


def reduce_filler_phrases(text: str) -> str:
    cleaned = text
    for pattern in ACADEMIC_FILLER_PATTERNS:
        cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)
    return re.sub(r"\s+([,.;:])", r"\1", cleaned)


def remove_old_parenthetical_citations(text: str) -> str:
    def replace_if_old(match: re.Match) -> str:
        citation = match.group(0)
        years = [int(year) for year in re.findall(r"(?:19\d{2}|20\d{2})", citation)]
        if years and any(year <= 2020 for year in years):
            return ""
        return citation

    cleaned = re.sub(r"\([^\)]*(?:19\d{2}|20\d{2})[^\)]*\)", replace_if_old, text)
    cleaned = re.sub(r"\s+([,.;:])", r"\1", cleaned)
    cleaned = re.sub(r"\(\s*\)", "", cleaned)
    return cleaned


def limit_dominant_keywords(text: str) -> tuple[str, dict[str, int]]:
    counts = {}
    result = text
    for keyword, limit in DOMINANT_KEYWORD_LIMITS.items():
        pattern = re.compile(re.escape(keyword), flags=re.IGNORECASE)
        matches = list(pattern.finditer(result))
        counts[keyword] = len(matches)
        if len(matches) <= limit:
            continue
        replacement = DOMINANT_KEYWORD_REPLACEMENTS[keyword]
        seen = 0

        def replace_after_limit(match: re.Match) -> str:
            nonlocal seen
            seen += 1
            if seen <= limit:
                return match.group(0)
            return replacement

        result = pattern.sub(replace_after_limit, result)
    return result, counts


def editorial_postprocess_section(section_text: str, section: dict, level: str) -> tuple[str, dict]:
    if section["key"] == "bibliografia":
        return cap_bibliography(section_text, level), {"removed_duplicate_paragraphs": 0, "keyword_counts": {}}

    processed = remove_old_parenthetical_citations(reduce_filler_phrases(section_text))
    blocks = re.split(r"\n\s*\n", processed)
    kept_blocks = []
    removed = 0
    for block in blocks:
        paragraph = block.strip()
        if not paragraph:
            continue
        if is_near_duplicate(paragraph, kept_blocks):
            removed += 1
            continue
        kept_blocks.append(paragraph)

    processed = "\n\n".join(kept_blocks).strip()
    processed, keyword_counts = limit_dominant_keywords(processed)
    report = {
        "removed_duplicate_paragraphs": removed,
        "keyword_counts": {key: value for key, value in keyword_counts.items() if value > DOMINANT_KEYWORD_LIMITS[key]},
    }
    return processed, report


def editorial_postprocess_document(content: str) -> tuple[str, dict]:
    blocks = re.split(r"\n\s*\n", content)
    kept = []
    seen_headings = set()
    removed_headings = 0
    for block in blocks:
        text = block.strip()
        if not text:
            continue
        normalized = normalize_heading(text)
        if text.isupper() and len(text) <= 90:
            if normalized in seen_headings:
                removed_headings += 1
                continue
            seen_headings.add(normalized)
        kept.append(text)

    processed = "\n\n".join(kept).strip()
    processed, keyword_counts = limit_dominant_keywords(processed)
    return processed, {
        "removed_duplicate_headings": removed_headings,
        "keyword_counts": {key: value for key, value in keyword_counts.items() if value > DOMINANT_KEYWORD_LIMITS[key]},
    }


def cap_bibliography(text: str, level: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    heading = []
    references = []
    for line in lines:
        if "bibliograf" in normalize_for_similarity(line):
            continue
        if re.search(r"\((20\d{2})[a-z]?\)", line):
            references.append(line)
    recent = []
    seen = set()
    for ref in references:
        years = [int(year) for year in re.findall(r"\((20\d{2})[a-z]?\)", ref)]
        if not years or any(year <= 2020 for year in years):
            continue
        key = normalize_for_similarity(ref)
        if key in seen:
            continue
        recent.append(ref)
        seen.add(key)
    capped = recent[:max_references_for_level(level)]
    return "\n\n".join((heading[:1] or []) + capped).strip() or text.strip()


def filter_recent_bibliography(items: List[str]) -> List[str]:
    recent_items = []
    for item in items:
        years = [int(year) for year in re.findall(r"\((\d{4})[a-z]?\)", item)]
        if years and all(year > 2020 for year in years):
            recent_items.append(item)
    return recent_items


def format_recent_bibliography(items: List[str]) -> str:
    recent_items = filter_recent_bibliography(items)
    if recent_items:
        return "\n".join(f"- {item}" for item in recent_items[:30])
    return "No se identifico bibliografia posterior a 2020 en el silabo. No inventes referencias especificas: usa solo fuentes institucionales, documentacion oficial o autores/obras que puedas reconocer como reales y verificables."


def min_references_for_level(level: str) -> int:
    return {
        "pregrado": 20,
        "especializacion": 12,
        "diplomado": 15,
        "maestria": 40,
    }[level]


def max_references_for_level(level: str) -> int:
    return {
        "pregrado": 30,
        "especializacion": 18,
        "diplomado": 25,
        "maestria": 55,
    }[level]


MAX_SECTION_EXPANSION_ATTEMPTS = 3


def build_expansion_prompt(plan: CoursePlan, topic: str, section: dict, current_text: str) -> str:
    return f"""
La siguiente seccion quedo demasiado corta para un documento academico de 20 a 30 paginas.

Tema: {topic}
Asignatura: {plan.asignatura}
Nivel academico: {plan.nivel}
Seccion: {section["title"]}
Minimo requerido para esta seccion: {section["min_words"]} palabras
Palabras actuales aproximadas: {word_count(current_text)}

Texto actual de la seccion:
{current_text[-6000:]}

Amplia esta misma seccion con contenido adicional sustantivo.
Reglas:
- No repitas literalmente lo ya escrito.
- No reformules argumentos ya presentes ni agregues parrafos de relleno para cumplir extension.
- No agregues encabezados principales nuevos.
- Mantén el mismo tono de guion editorial academico.
- Desarrolla unicamente contenido nuevo: ejemplos, analisis comparativo, implicaciones tecnicas, criterios de decision, tensiones profesionales o casos aplicados distintos a los ya usados.
- Evita repetir en exceso "contexto latinoamericano", "innovacion", "inmersion", "narrativa interactiva", "accesibilidad" y "sostenibilidad".
- Escribe solo la ampliacion que debe anexarse a esta seccion.
""".strip()


def build_bibliography_rewrite_prompt(plan: CoursePlan, topic: str, current_text: str) -> str:
    min_refs = min_references_for_level(plan.nivel)
    max_refs = max_references_for_level(plan.nivel)
    return f"""
La bibliografia generada no cumple el requisito institucional.

Tema: {topic}
Asignatura: {plan.asignatura}
Nivel academico: {plan.nivel}

Bibliografia actual:
{current_text[-6000:]}

Reescribe la bibliografia completa.
Reglas obligatorias:
- Escribe únicamente BIBLIOGRAFIA en formato APA 7.
- Incluye entre {min_refs} y {max_refs} referencias; no agregues bibliografia extensa para aparentar profundidad.
- Todas las referencias deben ser posteriores a 2020, es decir, de 2021 en adelante.
- No incluyas ninguna referencia de 2020, 2019, 2018 ni años anteriores.
- Usa solo fuentes reales, reconocibles y pertinentes al tema: bibliografia del silabo, documentacion oficial, organismos reconocidos, normas o autores/obras que puedas identificar como verificables.
- No inventes autores, titulos, revistas, editoriales, DOI, URL, volumenes, numeros ni paginas. Si no estas seguro de un dato, omitelo.
- Escribe una referencia por parrafo.
""".strip()


def expand_if_short(
    client: OpenAI,
    model: str,
    system_prompt: str,
    plan: CoursePlan,
    topic: str,
    section: dict,
    section_text: str,
    max_tokens: int,
    temperature: float,
) -> str:
    if section["key"] == "bibliografia":
        min_refs = min_references_for_level(plan.nivel)
        if recent_reference_count(section_text) >= min_refs and not has_old_references(section_text):
            return section_text
        print(f"    La bibliografia no cumple {min_refs} referencias posteriores a 2020; solicitando reescritura.")
        return generate_document(
            client=client,
            model=model,
            system_prompt=system_prompt,
            user_prompt=build_bibliography_rewrite_prompt(plan, topic, section_text),
            max_tokens=max_tokens,
            temperature=temperature,
        )
    elif word_count(section_text) >= section["min_words"]:
        return section_text

    combined = section_text
    for attempt in range(1, MAX_SECTION_EXPANSION_ATTEMPTS + 1):
        print(f"    La seccion quedo corta; solicitando ampliacion (intento {attempt}/{MAX_SECTION_EXPANSION_ATTEMPTS}).")
        try:
            expansion = generate_document(
                client=client,
                model=model,
                system_prompt=system_prompt,
                user_prompt=build_expansion_prompt(plan, topic, section, combined),
                max_tokens=max_tokens,
                temperature=temperature,
            )
        except Exception as exc:
            print(f"    ERROR en ampliacion: {type(exc).__name__}: {exc}")
            traceback.print_exc(file=sys.stderr)
            if attempt >= MAX_SECTION_EXPANSION_ATTEMPTS:
                print(
                    f"    ADVERTENCIA: no se pudo ampliar la seccion {section['key']} tras {MAX_SECTION_EXPANSION_ATTEMPTS} intentos; "
                    f"se continua con ~{word_count(combined)} palabras (minimo institucional {section['min_words']}).",
                )
                return combined
            continue
        combined = (combined.rstrip() + "\n\n" + strip_repeated_title(expansion, section["title"])).strip()
        if word_count(combined) >= section["min_words"]:
            return combined

    print(
        f"    ADVERTENCIA: tras {MAX_SECTION_EXPANSION_ATTEMPTS} ampliaciones la seccion {section['key']} sigue corta "
        f"({word_count(combined)} < {section['min_words']}); se continua con el texto disponible.",
    )
    return combined


def generate_long_document(
    client: OpenAI,
    model: str,
    system_prompt: str,
    plan: CoursePlan,
    syllabus_text: str,
    topic: str,
    topic_index: int,
    total_topics: int,
    max_tokens: int,
    temperature: float,
) -> str:
    header = (
        f"CONTENIDO: {topic}. ASIGNATURA: {plan.asignatura}. "
        f"PROGRAMA: {plan.programa}. NIVEL: {plan.nivel}. CICLO: {plan.ciclo}. SEMESTRE: {plan.semestre}."
    )
    blocks = [header]
    previous_context = ""
    used_main_titles = set()

    for section in get_section_plan(plan.nivel):
        print(f"  - Generando {section['key']} ({section['target_words']})")
        prompt = build_section_prompt(
            plan=plan,
            syllabus_text=syllabus_text,
            topic=topic,
            topic_index=topic_index,
            total_topics=total_topics,
            section=section,
            previous_context=previous_context,
        )
        section_text = generate_document(
            client=client,
            model=model,
            system_prompt=system_prompt,
            user_prompt=prompt,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        section_text = strip_repeated_title(section_text, section["title"])
        section_text = expand_if_short(
            client=client,
            model=model,
            system_prompt=system_prompt,
            plan=plan,
            topic=topic,
            section=section,
            section_text=section_text,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        section_text, editorial_report = editorial_postprocess_section(section_text, section, plan.nivel)
        if editorial_report["removed_duplicate_paragraphs"]:
            print(f"    Postproceso editorial: {editorial_report['removed_duplicate_paragraphs']} parrafos redundantes removidos.")
        if editorial_report["keyword_counts"]:
            dominant = ", ".join(f"{key}={value}" for key, value in editorial_report["keyword_counts"].items())
            print(f"    Postproceso editorial: keywords dominantes ajustadas ({dominant}).")

        if section["title"] not in used_main_titles:
            blocks.append(section["title"])
            used_main_titles.add(section["title"])
        blocks.append(section_text)
        previous_context = (previous_context + "\n\n" + section_text).strip()

    content, document_report = editorial_postprocess_document("\n\n".join(blocks).strip())
    if document_report["removed_duplicate_headings"]:
        print(f"  Postproceso editorial global: {document_report['removed_duplicate_headings']} subtitulos duplicados removidos.")
    if document_report["keyword_counts"]:
        dominant = ", ".join(f"{key}={value}" for key, value in document_report["keyword_counts"].items())
        print(f"  Postproceso editorial global: keywords dominantes ajustadas ({dominant}).")
    return content


def strip_repeated_title(text: str, title: str) -> str:
    lines = text.strip().splitlines()
    while lines and is_generated_wrapper_line(lines[0], title):
        lines.pop(0)
    return "\n".join(lines).strip()


def normalize_heading(value: str) -> str:
    value = clean_text(value).upper()
    replacements = {
        "Á": "A",
        "É": "E",
        "Í": "I",
        "Ó": "O",
        "Ú": "U",
        "Ü": "U",
        "Ñ": "N",
    }
    for original, replacement in replacements.items():
        value = value.replace(original, replacement)
    return value


def is_generated_wrapper_line(line: str, title: str) -> bool:
    normalized = normalize_heading(line)
    normalized_title = normalize_heading(title)
    return (
        normalized.startswith("CONTENIDO:")
        or normalized == normalized_title
        or normalized in {
            "INTRODUCCION",
            "EJES ARTICULADORES",
            "ENSAYOS DE PROFUNDIZACION",
            "CAPITULOS TEORICO-APLICADOS",
            "ANALISIS CRITICOS",
            "MODULOS TEMATICOS",
            "CASOS DE APLICACION PROFESIONAL",
            "CAPITULOS DE DESARROLLO",
            "DISCUSIONES TEORICAS",
            "CONCLUSIONES",
            "BIBLIOGRAFIA",
        }
    )


def save_docx(content: str, output_path: Path) -> None:
    doc = Document()
    for raw_block in re.split(r"\n\s*\n", content):
        block = raw_block.strip()
        if not block:
            continue
        if block.isupper() and len(block) <= 80:
            doc.add_heading(block, level=1)
        elif block.startswith("CONTENIDO:"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(block)
            run.bold = True
        else:
            for line in block.splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
    doc.save(output_path)


def write_plan_json(plan: CoursePlan, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / "plan_curso.json"
    path.write_text(json.dumps(asdict(plan), ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Genera 5 guiones académicos desde un sílabo DOCX o PDF.")
    parser.add_argument("--syllabus", required=True, help="Ruta al sílabo .docx o .pdf")
    parser.add_argument(
        "--nivel",
        default="auto",
        choices=["auto", *CATEGORIES.keys()],
        help="Nivel academico/prompt a usar. En auto intenta detectarlo desde el silabo.",
    )
    parser.add_argument("--subject", default="", help="Nombre de la materia, si se quiere forzar")
    parser.add_argument("--semester", default="", help="Semestre, si se quiere forzar")
    parser.add_argument("--topics", default="", help="Cinco temas separados por punto y coma o barra vertical, si se quieren forzar")
    parser.add_argument("--output-dir", default="outputs", help="Carpeta de salida")
    parser.add_argument("--prompt", default="", help="Ruta a un prompt maestro personalizado. Si se omite, usa prompts/<nivel>.md")
    parser.add_argument("--model", default=os.getenv("OPENAI_MODEL", "gpt-4o"), help="Modelo OpenAI")
    parser.add_argument("--max-tokens", type=int, default=4500, help="Maximo de tokens por seccion generada")
    parser.add_argument("--temperature", type=float, default=0.65, help="Creatividad de generación")
    parser.add_argument("--dry-run", action="store_true", help="Solo analiza el sílabo y muestra el plan")
    return parser.parse_args()


def main() -> None:
    if load_dotenv:
        load_dotenv()

    args = parse_args()
    syllabus_path = Path(args.syllabus)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    syllabus_text = extract_input_text(syllabus_path)
    plan = extract_course_plan(syllabus_path, args.subject, args.semester, args.topics)
    plan.nivel = detect_academic_level(plan, syllabus_text, args.nivel)
    prompt_path = resolve_prompt_path(plan.nivel, args.prompt)
    plan_path = write_plan_json(plan, output_dir)

    print("Plan detectado:")
    print(json.dumps(asdict(plan), ensure_ascii=False, indent=2))
    print(f"\nNivel seleccionado: {plan.nivel}")
    print(f"Prompt seleccionado: {prompt_path}")
    print(f"\nPlan guardado en: {plan_path}")

    if args.dry_run:
        print("\nDry-run activo. No se llamó a la API.")
        return

    if OpenAI is None:
        raise RuntimeError("Falta instalar el paquete openai. Ejecuta: pip install -r requirements.txt")
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("Falta OPENAI_API_KEY en variables de entorno.")

    client = OpenAI()
    system_prompt = load_system_prompt(prompt_path)

    for index, topic in enumerate(plan.temas, start=1):
        print(f"\nGenerando documento {index}/{len(plan.temas)}: {topic}")
        content = generate_long_document(
            client=client,
            model=args.model,
            system_prompt=system_prompt,
            plan=plan,
            syllabus_text=syllabus_text,
            topic=topic,
            topic_index=index,
            total_topics=len(plan.temas),
            max_tokens=args.max_tokens,
            temperature=args.temperature,
        )
        output_path = output_dir / f"G{index}_{slugify(topic)}.docx"
        save_docx(content, output_path)
        words = word_count(content)
        estimated_pages = round(words / 450, 1)
        print(f"Guardado: {output_path} ({words} palabras aprox.; {estimated_pages} paginas estimadas)")
        try:
            from automation_engine.incremental_drive_upload import upload_package_file_if_configured

            upload_package_file_if_configured(
                output_path,
                f"PAQUETE_ACADEMICO/CONTENIDOS/G{index}.docx",
            )
        except Exception as sync_exc:
            print(f"Drive incremental: aviso (no detiene generación): {sync_exc}")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        sys.stderr.write("\nInterrumpido por el usuario.\n")
        sys.stderr.flush()
        raise SystemExit(130)
    except Exception as exc:
        traceback.print_exc(file=sys.stderr)
        sys.stderr.write(f"\n[FATAL] {type(exc).__name__}: {exc}\n")
        sys.stderr.flush()
        raise SystemExit(1) from exc
